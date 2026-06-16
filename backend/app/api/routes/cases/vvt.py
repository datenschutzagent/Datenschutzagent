"""Case VVT normalization (Art. 30): canonical mapping plus CSV/DOCX export."""

import csv
import io
import logging
from datetime import UTC, datetime
from typing import Literal
from uuid import UUID

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.database import get_db
from app.models import CaseModel
from app.models.schemas import (
    VVTFieldResponse,
    VVTNormalizationResponse,
)
from app.services.org_profile_loader import get_vvt_field_names
from app.services.query_helpers import case_relations
from app.services.vvt_service import normalize_vvt

logger = logging.getLogger(__name__)

router = APIRouter()


def _build_vvt_export_docx(doc_name: str, source_template: str, fields: list) -> bytes:
    """Build DOCX with VVT normalization (Ziel-Template): document name, template, table of fields."""
    doc = DocxDocument()
    doc.add_heading("VVT-Normalisierung (Ziel-Template)", 0)
    doc.add_paragraph()
    doc.add_paragraph(f"Dokument: {doc_name}")
    doc.add_paragraph(f"Erkanntes Template: {source_template or '—'}")
    doc.add_paragraph()
    table = doc.add_table(rows=1 + len(fields), cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Feldname"
    hdr[1].text = "Status"
    hdr[2].text = "Kanonischer Wert"
    hdr[3].text = "Nachweis"
    hdr[4].text = "Hinweis"
    for i, f in enumerate(fields):
        row = table.rows[i + 1].cells
        row[0].text = (f.field_name or "").replace("\r\n", " ").replace("\n", " ")
        row[1].text = (f.status or "").replace("\r\n", " ").replace("\n", " ")
        row[2].text = (f.canonical_value or "").replace("\r\n", " ").replace("\n", " ")
        row[3].text = (f.evidence or "").replace("\r\n", " ").replace("\n", " ")
        row[4].text = (f.finding or "").replace("\r\n", " ").replace("\n", " ")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/{case_id}/vvt-normalization/export", response_class=Response)
async def get_vvt_normalization_export(
    case_id: UUID,
    document_id: UUID | None = Query(None, alias="document_id"),
    format: Literal["csv", "docx"] = Query("csv", alias="format"),
    db: AsyncSession = Depends(get_db),
):
    """
    Export VVT normalization. Query: format=csv (default) or format=docx.
    Uses first VVT document if document_id not given.
    CSV: text/csv. DOCX: VVT Ziel-Template (document name, template, fields table).
    """
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(*case_relations(findings=False))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    vvt_docs = [d for d in case.documents if d.type == "vvt"]
    if not vvt_docs:
        raise HTTPException(status_code=404, detail="No VVT document in this case")

    if document_id is not None:
        doc = next((d for d in vvt_docs if d.id == document_id), None)
        if not doc:
            raise HTTPException(
                status_code=404,
                detail="Document not found or is not a VVT document of this case",
            )
    else:
        doc = vvt_docs[0]

    raw_text = doc.content or ""
    if not raw_text.strip():
        raise HTTPException(
            status_code=404,
            detail="VVT-Dokument hat keinen extrahierten Inhalt",
        )

    case_lang = getattr(case, "language", None)
    vvt_fields = get_vvt_field_names(settings)
    extraction = await normalize_vvt(
        raw_text, language=case_lang, field_names=vvt_fields
    )
    date_str = datetime.now(UTC).strftime("%Y-%m-%d")

    if (format or "csv").lower() == "docx":
        content = _build_vvt_export_docx(
            doc.name, extraction.source_template or "", extraction.fields
        )
        filename = f"VVT-Ziel-{case_id}-{date_str}.docx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(
        [
            "document_name",
            "source_template",
            "field_name",
            "status",
            "canonical_value",
            "evidence",
            "finding",
        ]
    )
    for f in extraction.fields:
        writer.writerow(
            [
                doc.name,
                extraction.source_template or "",
                f.field_name,
                f.status,
                (f.canonical_value or "").replace("\r\n", " ").replace("\n", " "),
                (f.evidence or "").replace("\r\n", " ").replace("\n", " "),
                (f.finding or "").replace("\r\n", " ").replace("\n", " "),
            ]
        )
    csv_content = buf.getvalue()
    filename = f"VVT-Export-{case_id}-{date_str}.csv"
    body = "\ufeff" + csv_content  # UTF-8 BOM for Excel
    return Response(
        content=body.encode("utf-8"),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{case_id}/vvt-normalization", response_model=VVTNormalizationResponse)
async def get_vvt_normalization(
    case_id: UUID,
    document_id: UUID | None = None,
    db: AsyncSession = Depends(get_db),
):
    """
    Get VVT normalization for the case. Uses the first VVT document if document_id is not given.
    Returns canonical VVT fields and template detection (LLM-based).
    """
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(*case_relations(findings=False))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    vvt_docs = [d for d in case.documents if d.type == "vvt"]
    if not vvt_docs:
        return VVTNormalizationResponse(
            document_id=None,
            document_name="",
            source_template="",
            fields=[],
        )

    if document_id is not None:
        doc = next((d for d in vvt_docs if d.id == document_id), None)
        if not doc:
            raise HTTPException(
                status_code=404,
                detail="Document not found or is not a VVT document of this case",
            )
    else:
        doc = vvt_docs[0]

    raw_text = doc.content or ""
    if not raw_text.strip():
        return VVTNormalizationResponse(
            document_id=doc.id,
            document_name=doc.name,
            source_template="",
            fields=[],
        )

    case_lang = getattr(case, "language", None)
    vvt_fields = get_vvt_field_names(settings)
    extraction = await normalize_vvt(
        raw_text, language=case_lang, field_names=vvt_fields
    )

    fields = [
        VVTFieldResponse(
            field_name=f.field_name,
            required=True,
            status=f.status,
            source_template=extraction.source_template,
            canonical_value=f.canonical_value,
            evidence=f.evidence,
            finding=f.finding,
        )
        for f in extraction.fields
    ]
    return VVTNormalizationResponse(
        document_id=doc.id,
        document_name=doc.name,
        source_template=extraction.source_template or "Unbekannt",
        fields=fields,
    )
