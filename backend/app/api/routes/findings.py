"""Findings API: list, update, bulk-update, CSV export, DOCX export, comments."""
import csv
import io
from datetime import datetime, timezone
from typing import Annotated
from uuid import UUID

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from sqlalchemy import func, select, text
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.auth import require_roles
from app.database import get_db
from app.models.db import (
    ActivityLogModel,
    CaseModel,
    DocumentModel,
    FindingChatMessageModel,
    FindingCommentModel,
    FindingModel,
    UserModel,
)
from app.models.schemas import (
    FindingBulkDelete,
    FindingBulkUpdate,
    FindingChatMessageCreate,
    FindingChatMessageResponse,
    FindingCommentCreate,
    FindingCommentResponse,
    FindingListResponse,
    FindingResponse,
    FindingStatsResponse,
    FindingTrendItem,
    FindingsByDepartment,
    TopFailingCheck,
    FindingUpdate,
)

router = APIRouter()

EXPORT_MAX = 5000


@router.get("/stats", response_model=FindingStatsResponse, summary="Befund-Statistiken abrufen")
async def get_findings_stats(
    db: AsyncSession = Depends(get_db),
    _user=require_roles("viewer", "editor", "admin"),
):
    """Aggregierte Befund-Statistiken: nach Schweregrad, Kategorie, Abteilung, Top-Checks und Zeitverlauf (6 Monate).
    Alle Teilabfragen werden in einem einzigen DB-Round-Trip via CTEs ausgeführt."""

    stats_result = await db.execute(text("""
        WITH by_severity AS (
            SELECT severity, COUNT(*) AS cnt
            FROM findings
            WHERE status = 'open'
            GROUP BY severity
        ),
        by_category AS (
            SELECT category, COUNT(*) AS cnt
            FROM findings
            WHERE status = 'open'
            GROUP BY category
            ORDER BY cnt DESC
            LIMIT 20
        ),
        by_department AS (
            SELECT c.department,
                   SUM(CASE WHEN f.severity = 'critical' THEN 1 ELSE 0 END) AS critical,
                   SUM(CASE WHEN f.severity = 'high'     THEN 1 ELSE 0 END) AS high,
                   SUM(CASE WHEN f.severity = 'medium'   THEN 1 ELSE 0 END) AS medium,
                   SUM(CASE WHEN f.severity = 'low'      THEN 1 ELSE 0 END) AS low,
                   SUM(CASE WHEN f.severity = 'info'     THEN 1 ELSE 0 END) AS info,
                   COUNT(*) AS total
            FROM findings f
            JOIN cases c ON c.id = f.case_id
            WHERE f.status = 'open'
            GROUP BY c.department
            ORDER BY total DESC
            LIMIT 20
        ),
        top_checks AS (
            SELECT check_name, category,
                   COUNT(*) AS total,
                   SUM(CASE WHEN severity = 'critical' THEN 1 ELSE 0 END) AS critical,
                   SUM(CASE WHEN severity = 'high'     THEN 1 ELSE 0 END) AS high,
                   SUM(CASE WHEN severity = 'medium'   THEN 1 ELSE 0 END) AS medium,
                   SUM(CASE WHEN severity = 'low'      THEN 1 ELSE 0 END) AS low,
                   SUM(CASE WHEN severity = 'info'     THEN 1 ELSE 0 END) AS info
            FROM findings
            WHERE status = 'open'
            GROUP BY check_name, category
            ORDER BY total DESC
            LIMIT 10
        ),
        trend AS (
            SELECT TO_CHAR(DATE_TRUNC('month', created_at), 'YYYY-MM') AS month,
                   SUM(CASE WHEN severity = 'critical' THEN 1 ELSE 0 END) AS critical,
                   SUM(CASE WHEN severity = 'high'     THEN 1 ELSE 0 END) AS high,
                   SUM(CASE WHEN severity = 'medium'   THEN 1 ELSE 0 END) AS medium,
                   SUM(CASE WHEN severity = 'low'      THEN 1 ELSE 0 END) AS low,
                   SUM(CASE WHEN severity = 'info'     THEN 1 ELSE 0 END) AS info
            FROM findings
            WHERE created_at >= NOW() - INTERVAL '6 months'
            GROUP BY DATE_TRUNC('month', created_at)
            ORDER BY DATE_TRUNC('month', created_at) ASC
        )
        SELECT
            'severity'   AS qname, severity   AS key1, NULL AS key2, cnt AS n,  0 AS c2, 0 AS c3, 0 AS c4, 0 AS c5, 0 AS c6  FROM by_severity
        UNION ALL
        SELECT 'category', category, NULL, cnt, 0, 0, 0, 0, 0                                                                   FROM by_category
        UNION ALL
        SELECT 'department', department, NULL, total, critical, high, medium, low, info                                         FROM by_department
        UNION ALL
        SELECT 'top_check', check_name, category, total, critical, high, medium, low, info                                      FROM top_checks
        UNION ALL
        SELECT 'trend', month, NULL, 0, critical, high, medium, low, info                                                       FROM trend
    """))
    rows = stats_result.fetchall()

    by_severity: dict[str, int] = {}
    by_category: dict[str, int] = {}
    by_department: list[FindingsByDepartment] = []
    top_failing_checks: list[TopFailingCheck] = []
    trend: list[FindingTrendItem] = []

    for row in rows:
        qname, key1, key2, n, c2, c3, c4, c5, c6 = row
        if qname == "severity":
            by_severity[key1] = n
        elif qname == "category":
            by_category[key1] = n
        elif qname == "department":
            by_department.append(FindingsByDepartment(
                department=key1, total=n, critical=c2, high=c3, medium=c4, low=c5, info=c6,
            ))
        elif qname == "top_check":
            top_failing_checks.append(TopFailingCheck(
                check_name=key1,
                category=key2 or "",
                count=n,
                severity_breakdown={"critical": c2, "high": c3, "medium": c4, "low": c5, "info": c6},
            ))
        elif qname == "trend":
            trend.append(FindingTrendItem(
                month=key1, critical=c2, high=c3, medium=c4, low=c5, info=c6,
            ))

    return FindingStatsResponse(
        by_severity=by_severity,
        by_category=by_category,
        by_department=by_department,
        top_failing_checks=top_failing_checks,
        trend=trend,
    )


@router.get("", response_model=FindingListResponse, summary="Findings auflisten")
async def list_findings(
    case_id: Annotated[UUID | None, Query()] = None,
    q: Annotated[str | None, Query(description="Volltext-Suche in Check-Name und Beschreibung")] = None,
    severity: Annotated[str | None, Query()] = None,
    status: Annotated[str | None, Query()] = None,
    category: Annotated[str | None, Query()] = None,
    source_strategy: Annotated[str | None, Query(description="Filter by check source: full_text or rag")] = None,
    has_due_date: Annotated[bool | None, Query(description="True = only findings with a due date set")] = None,
    limit: Annotated[int, Query(ge=1, le=500)] = 100,
    offset: Annotated[int, Query(ge=0)] = 0,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("viewer", "editor", "admin"),
):
    """List findings with optional filters. Returns paginated results."""
    from sqlalchemy import or_
    query = select(FindingModel)
    if case_id is not None:
        query = query.where(FindingModel.case_id == case_id)
    if q:
        like = f"%{q}%"
        query = query.where(
            or_(
                FindingModel.check_name.ilike(like),
                FindingModel.description.ilike(like),
            )
        )
    if severity is not None:
        query = query.where(FindingModel.severity == severity)
    if status is not None:
        query = query.where(FindingModel.status == status)
    if category is not None:
        query = query.where(FindingModel.category == category)
    if source_strategy is not None:
        query = query.where(FindingModel.source_strategy == source_strategy)
    if has_due_date is True:
        query = query.where(FindingModel.due_date != None)  # noqa: E711
    elif has_due_date is False:
        query = query.where(FindingModel.due_date == None)  # noqa: E711

    count_q = select(func.count()).select_from(query.subquery())
    total_result = await db.execute(count_q)
    total = total_result.scalar_one()

    result = await db.execute(
        query.offset(offset).limit(limit).options(selectinload(FindingModel.case))
    )
    findings = result.scalars().all()

    return FindingListResponse(
        items=[FindingResponse.model_validate(f) for f in findings],
        total=total,
    )


@router.patch("/bulk-update")
async def bulk_update_findings(
    body: FindingBulkUpdate,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Set the same status on multiple findings at once."""
    if not body.finding_ids:
        return {"updated": 0}

    result = await db.execute(
        select(FindingModel).where(FindingModel.id.in_(body.finding_ids))
    )
    findings = result.scalars().all()

    updated = 0
    for finding in findings:
        if finding.status != body.status:
            old_status = finding.status
            finding.status = body.status
            activity = ActivityLogModel(
                case_id=finding.case_id,
                event_type="finding_status_updated",
                payload={
                    "finding_id": str(finding.id),
                    "old_status": old_status,
                    "new_status": body.status,
                    "bulk": True,
                },
            )
            db.add(activity)
        updated += 1

    await db.flush()
    return {"updated": updated}


@router.delete("/bulk-delete", status_code=200)
async def bulk_delete_findings(
    body: FindingBulkDelete,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Mehrere Befunde auf einmal löschen. Gibt die Anzahl gelöschter Befunde zurück."""
    if not body.finding_ids:
        return {"deleted": 0}

    result = await db.execute(
        select(FindingModel).where(FindingModel.id.in_(body.finding_ids))
    )
    findings = result.scalars().all()
    for finding in findings:
        await db.delete(finding)
    await db.flush()
    return {"deleted": len(findings)}


def _build_findings_docx(case_title: str, findings: list, docs_by_id: dict) -> bytes:
    """Build DOCX with structured findings report: title page, summary table, per-finding sections."""
    severity_labels = {
        "critical": "Kritisch", "high": "Hoch", "medium": "Mittel", "low": "Niedrig", "info": "Info",
    }
    status_labels = {
        "open": "Offen", "accepted": "Akzeptiert", "overruled": "Überfahren", "fixed": "Behoben",
    }

    doc = DocxDocument()
    doc.add_heading("Befunde-Bericht", 0)
    doc.add_paragraph(f"Vorgang: {case_title}")
    doc.add_paragraph(f"Erstellt: {datetime.now(timezone.utc).strftime('%d.%m.%Y')}")
    doc.add_paragraph()

    # Summary table
    severity_counts = {s: sum(1 for f in findings if f.severity == s) for s in severity_labels}
    status_counts = {s: sum(1 for f in findings if f.status == s) for s in status_labels}
    doc.add_heading("Zusammenfassung", level=1)
    summary_table = doc.add_table(rows=1 + len(severity_labels), cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Schweregrad"
    summary_table.rows[0].cells[1].text = "Anzahl"
    for i, (sev, label) in enumerate(severity_labels.items()):
        row = summary_table.rows[i + 1]
        row.cells[0].text = label
        row.cells[1].text = str(severity_counts.get(sev, 0))
    doc.add_paragraph()

    # Per-finding sections
    doc.add_heading("Einzelbefunde", level=1)
    for f in findings:
        doc.add_heading(f.check_name, level=2)
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = "Table Grid"
        labels_rows = [
            ("Schweregrad", severity_labels.get(f.severity, f.severity)),
            ("Status", status_labels.get(f.status, f.status)),
            ("Kategorie", f.category),
            ("Dokument", docs_by_id.get(f.document_id, "Vorgangsbezogen") if f.document_id else "Vorgangsbezogen"),
            ("Strategie", f.source_strategy or ""),
        ]
        for i, (label, value) in enumerate(labels_rows):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
        doc.add_paragraph()
        doc.add_paragraph("Beschreibung:").bold = True
        doc.add_paragraph(f.description)
        if f.evidence:
            doc.add_paragraph("Nachweise:").bold = True
            for ev in f.evidence:
                doc.add_paragraph(f"• {ev}")
        if f.recommendation:
            doc.add_paragraph("Empfehlung:").bold = True
            doc.add_paragraph(f.recommendation)
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/export", response_class=Response)
async def export_findings(
    case_id: Annotated[UUID | None, Query()] = None,
    severity: Annotated[str | None, Query()] = None,
    status: Annotated[str | None, Query()] = None,
    category: Annotated[str | None, Query()] = None,
    format: Annotated[str, Query()] = "csv",
    db: AsyncSession = Depends(get_db),
    _user=require_roles("viewer", "editor", "admin"),
):
    """Export findings as CSV (default) or DOCX. If case_id is omitted, exports all findings (org-wide)."""
    q = select(FindingModel)
    if case_id is not None:
        q = q.where(FindingModel.case_id == case_id)
    if severity is not None:
        q = q.where(FindingModel.severity == severity)
    if status is not None:
        q = q.where(FindingModel.status == status)
    if category is not None:
        q = q.where(FindingModel.category == category)

    # Count total before applying limit so we can signal truncation
    count_result = await db.execute(select(func.count()).select_from(q.subquery()))
    total_count = count_result.scalar_one()

    findings_result = await db.execute(q.limit(EXPORT_MAX))
    findings = findings_result.scalars().all()
    is_truncated = total_count > EXPORT_MAX

    case_title = "Alle Vorgänge"
    if case_id is not None:
        case_result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
        case = case_result.scalar_one_or_none()
        if case is None:
            raise HTTPException(status_code=404, detail="Case not found")
        case_title = case.title

    doc_ids = {f.document_id for f in findings if f.document_id is not None}
    docs_by_id: dict[UUID, str] = {}
    if doc_ids:
        docs_result = await db.execute(
            select(DocumentModel).where(DocumentModel.id.in_(doc_ids))
        )
        docs_by_id = {d.id: d.name for d in docs_result.scalars().all()}

    # Build case title map for CSV (when exporting org-wide)
    case_ids = {f.case_id for f in findings}
    cases_by_id: dict[UUID, str] = {}
    if case_ids:
        cases_result = await db.execute(select(CaseModel).where(CaseModel.id.in_(case_ids)))
        cases_by_id = {c.id: c.title for c in cases_result.scalars().all()}

    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    slug = case_title.replace("/", "-").replace("\\", "-")[:50]

    export_headers = {
        "X-Total-Count": str(total_count),
        "X-Truncated": str(is_truncated).lower(),
    }

    if format.lower() == "docx":
        content_bytes = _build_findings_docx(case_title, findings, docs_by_id)
        filename = f"Befunde-{slug}-{date_str}.docx"
        return Response(
            content=content_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"', **export_headers},
        )

    severity_labels = {
        "critical": "Kritisch",
        "high": "Hoch",
        "medium": "Mittel",
        "low": "Niedrig",
        "info": "Info",
    }
    status_labels = {
        "open": "Offen",
        "accepted": "Akzeptiert",
        "overruled": "Überfahren",
        "fixed": "Behoben",
    }

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "ID",
        "Vorgang",
        "Checkname",
        "Kategorie",
        "Schweregrad",
        "Status",
        "Beschreibung",
        "Empfehlung",
        "Dokument",
        "Strategie",
    ])
    for f in findings:
        writer.writerow([
            str(f.id),
            cases_by_id.get(f.case_id, str(f.case_id)),
            f.check_name,
            f.category,
            severity_labels.get(f.severity, f.severity),
            status_labels.get(f.status, f.status),
            f.description,
            f.recommendation,
            docs_by_id.get(f.document_id, "") if f.document_id else "",
            f.source_strategy or "",
        ])

    filename = f"Befunde-{slug}-{date_str}.csv"

    # UTF-8 BOM for Excel compatibility (consistent with VVT export)
    body = "\ufeff" + buf.getvalue()
    return Response(
        content=body.encode("utf-8"),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"', **export_headers},
    )


@router.patch("/{finding_id}", response_model=FindingResponse)
async def update_finding(
    finding_id: UUID,
    body: FindingUpdate,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Update a finding (status and/or due_date)."""
    result = await db.execute(select(FindingModel).where(FindingModel.id == finding_id))
    finding = result.scalar_one_or_none()
    if not finding:
        raise HTTPException(status_code=404, detail="Finding not found")
    old_status = finding.status
    finding.status = body.status
    if body.due_date is not None or "due_date" in body.model_fields_set:
        finding.due_date = body.due_date
    await db.flush()
    if old_status != body.status:
        actor = _user.display_name if isinstance(_user, UserModel) else "System"
        activity = ActivityLogModel(
            case_id=finding.case_id,
            event_type="finding_status_updated",
            payload={
                "finding_id": str(finding_id),
                "check_name": finding.check_name,
                "old_status": old_status,
                "new_status": body.status,
                "actor": actor,
            },
        )
        db.add(activity)
    await db.refresh(finding)
    return FindingResponse.model_validate(finding)


# --- Finding Comments ---

@router.get("/{finding_id}/comments", response_model=list[FindingCommentResponse])
async def list_finding_comments(
    finding_id: UUID,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("viewer", "editor", "admin"),
):
    """List comments for a finding, sorted oldest-first."""
    result = await db.execute(select(FindingModel).where(FindingModel.id == finding_id))
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Finding not found")
    comments_result = await db.execute(
        select(FindingCommentModel)
        .where(FindingCommentModel.finding_id == finding_id)
        .order_by(FindingCommentModel.created_at.asc())
    )
    return [FindingCommentResponse.model_validate(c) for c in comments_result.scalars().all()]


@router.post("/{finding_id}/comments", response_model=FindingCommentResponse, status_code=201)
async def create_finding_comment(
    finding_id: UUID,
    body: FindingCommentCreate,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Add a comment to a finding (audit trail for status decisions)."""
    result = await db.execute(select(FindingModel).where(FindingModel.id == finding_id))
    finding = result.scalar_one_or_none()
    if not finding:
        raise HTTPException(status_code=404, detail="Finding not found")

    display_name = "System"
    user_id = None
    if isinstance(_user, UserModel):
        display_name = _user.display_name or "System"
        user_id = _user.id

    comment = FindingCommentModel(
        finding_id=finding_id,
        case_id=finding.case_id,
        author=display_name,
        user_id=user_id,
        text=body.text,
    )
    db.add(comment)
    activity = ActivityLogModel(
        case_id=finding.case_id,
        event_type="comment_added",
        payload={"finding_id": str(finding_id), "author": display_name},
    )
    db.add(activity)
    await db.flush()
    await db.refresh(comment)
    return FindingCommentResponse.model_validate(comment)


# --- LLM-Befund-Chat (KI-Assistent) ---

@router.get("/{finding_id}/chat", response_model=list[FindingChatMessageResponse], summary="Chat-Verlauf abrufen")
async def get_finding_chat(
    finding_id: UUID,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("viewer", "editor", "admin"),
):
    """Gibt alle Chat-Nachrichten für einen Befund zurück (älteste zuerst)."""
    result = await db.execute(select(FindingModel).where(FindingModel.id == finding_id))
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Finding not found")
    chat_result = await db.execute(
        select(FindingChatMessageModel)
        .where(FindingChatMessageModel.finding_id == finding_id)
        .order_by(FindingChatMessageModel.created_at.asc())
    )
    return [
        FindingChatMessageResponse.model_validate(m)
        for m in chat_result.scalars().all()
    ]


@router.post("/{finding_id}/chat", response_model=FindingChatMessageResponse, status_code=201, summary="Nachricht an KI-Assistent senden")
async def send_finding_chat_message(
    finding_id: UUID,
    body: FindingChatMessageCreate,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Sendet eine Nutzernachricht und erhält eine KI-Antwort im Befund-Kontext."""
    from app.services.finding_chat_service import chat_with_finding
    result = await db.execute(select(FindingModel).where(FindingModel.id == finding_id))
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Finding not found")

    try:
        response_text = await chat_with_finding(finding_id, body.content, db)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"LLM-Fehler: {exc}")

    # Die letzte Assistent-Nachricht zurückgeben
    last_result = await db.execute(
        select(FindingChatMessageModel)
        .where(
            FindingChatMessageModel.finding_id == finding_id,
            FindingChatMessageModel.role == "assistant",
        )
        .order_by(FindingChatMessageModel.created_at.desc())
        .limit(1)
    )
    last_msg = last_result.scalar_one_or_none()
    if not last_msg:
        raise HTTPException(status_code=500, detail="Chat message not saved")
    return FindingChatMessageResponse.model_validate(last_msg)
