"""Case management CRUD API."""
import asyncio
import csv
import io
import json as _json
import logging
import re
from datetime import datetime, timezone
from typing import Literal
from uuid import UUID

from docx import Document as DocxDocument

from app.constants import CaseStatus, DocumentExtractionStatus, FindingStatus, JobStatus

logger = logging.getLogger(__name__)

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import JSONResponse, Response, StreamingResponse
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.database import get_db
from app.services.risk_config_loader import get_risk_config
from app.models import CaseCreate, CaseListResponse, CaseModel, CaseResponse, CaseUpdate
from app.models.db import (
    ActivityLogModel,
    DocumentModel,
    DSBReportJobModel,
    PlaybookModel,
    RunChecksJobModel,
    UserModel,
)
from app.models.schemas import (
    ActivityResponse,
    AnnotatedDocumentListItem,
    CaseCloneRequest,
    CaseRiskScoreHistoryItem,
    CaseRiskScoreResponse,
    CaseSimilarityResult,
    DSBReportResponse,
    RunChecksRequest,
    VVTFieldResponse,
    VVTNormalizationResponse,
)
from app.services.annotated_document_service import (
    build_annotated_docx,
    build_annotated_pdf,
    list_annotatable_documents,
)
from app.services.run_checks_service import run_checks_impl
from app.celery_app import build_dsb_report_task, run_playbook_checks
from app.core.request_id import get_request_id
from app.services.weaviate_service import delete_chunks_by_case_id
from app.services.dsb_report_service import (
    build_dsb_report,
    compute_stale,
    get_last_run_checks_at,
    get_latest_report,
    render_report_markdown,
    save_report,
)
from app.services.dsb_report_service import _payload_to_report
from app.services.vvt_service import normalize_vvt
from app.services.org_profile_loader import get_vvt_field_names
from app.core.auth import require_roles
from app.core.rate_limit import limiter

router = APIRouter()


# ---- Running-checks endpoint (must be before /{case_id} routes) ----------

@router.get("/running-checks", summary="Alle laufenden Playbook-Prüfungen")
async def list_running_checks(db: AsyncSession = Depends(get_db)):
    """Return all currently running playbook check jobs across all cases.

    Used by the frontend to display a global status banner showing active checks.
    """
    result = await db.execute(
        select(RunChecksJobModel, CaseModel.title)
        .join(CaseModel, RunChecksJobModel.case_id == CaseModel.id)
        .where(RunChecksJobModel.status == JobStatus.RUNNING)
        .order_by(RunChecksJobModel.created_at.desc())
    )
    return [
        {
            "job_id": str(job.id),
            "case_id": str(job.case_id),
            "case_title": title,
            "playbook_name": job.playbook_name,
            "status": job.status,
            "checks_total": job.checks_total,
            "checks_done": job.checks_done,
            "created_at": job.created_at.isoformat() if job.created_at else None,
        }
        for job, title in result.all()
    ]


_CASE_SORT_COLUMNS = {
    "updated_at": CaseModel.updated_at,
    "created_at": CaseModel.created_at,
    "deadline": CaseModel.deadline,
    "title": CaseModel.title,
    "status": CaseModel.status,
}


@router.get("", response_model=CaseListResponse, summary="Vorgänge auflisten")
async def list_cases(
    skip: int = Query(default=0, ge=0),
    limit: int = Query(default=100, ge=1, le=500),
    q: str | None = Query(default=None, description="Volltext-Suche in Titel, Abteilung, Vorgangstyp, Ersteller, Zugewiesenem"),
    status: str | None = Query(default=None, description="Filtert nach Case-Status (z. B. intake, in_review, completed)"),
    department: str | None = Query(default=None, description="Filtert nach Abteilung (exakter Name)"),
    assignee: str | None = Query(default=None, description="Filtert nach Zugewiesenem (enthält)"),
    created_by: str | None = Query(default=None, description="Filtert nach Ersteller (enthält)"),
    has_open_findings: bool | None = Query(default=None, description="True = nur Vorgänge mit mindestens einem offenen Befund"),
    deadline_overdue: bool | None = Query(default=None, description="True = nur überfällige Vorgänge (Frist < heute, Status nicht completed)"),
    include_archived: bool = Query(default=False, description="True = auch archivierte Vorgänge zurückgeben (Standard: nur nicht-archivierte)"),
    sort_by: str = Query(default="updated_at", description="Sortierfeld: updated_at (Standard), created_at, deadline, title, status"),
    order: str = Query(default="desc", description="Sortierrichtung: desc (Standard) oder asc"),
    db: AsyncSession = Depends(get_db),
):
    """List cases with optional pagination and server-side filtering. Returns items and total count."""
    from datetime import date as date_type
    from sqlalchemy import exists
    base_q = select(CaseModel)
    if q:
        like = f"%{q}%"
        from sqlalchemy import or_
        base_q = base_q.where(
            or_(
                CaseModel.title.ilike(like),
                CaseModel.department.ilike(like),
                CaseModel.case_type.ilike(like),
                CaseModel.created_by.ilike(like),
                CaseModel.assignee.ilike(like),
            )
        )
    if status:
        base_q = base_q.where(CaseModel.status == status)
    if department:
        base_q = base_q.where(CaseModel.department == department)
    if assignee:
        base_q = base_q.where(CaseModel.assignee.ilike(f"%{assignee}%"))
    if created_by:
        base_q = base_q.where(CaseModel.created_by.ilike(f"%{created_by}%"))
    if has_open_findings is True:
        from app.models.db import FindingModel
        base_q = base_q.where(
            exists(select(FindingModel.id).where(
                FindingModel.case_id == CaseModel.id,
                FindingModel.status == FindingStatus.OPEN,
            ))
        )
    elif has_open_findings is False:
        from app.models.db import FindingModel
        base_q = base_q.where(
            ~exists(select(FindingModel.id).where(
                FindingModel.case_id == CaseModel.id,
                FindingModel.status == FindingStatus.OPEN,
            ))
        )
    if deadline_overdue is True:
        today = date_type.today()
        base_q = base_q.where(CaseModel.deadline < today, CaseModel.status != CaseStatus.COMPLETED)
    elif deadline_overdue is False:
        today = date_type.today()
        base_q = base_q.where((CaseModel.deadline >= today) | (CaseModel.deadline == None))  # noqa: E711
    if not include_archived:
        base_q = base_q.where(CaseModel.archived_at == None)  # noqa: E711

    count_result = await db.execute(select(func.count()).select_from(base_q.subquery()))
    total = count_result.scalar_one()

    sort_col = _CASE_SORT_COLUMNS.get(sort_by, CaseModel.updated_at)
    sort_expr = sort_col.asc() if order == "asc" else sort_col.desc()
    # NULL-Werte (z. B. deadline) werden bei DESC ans Ende gestellt (NULLS LAST)
    from sqlalchemy import nulls_last, nulls_first
    sort_expr = nulls_last(sort_expr) if order == "desc" else nulls_first(sort_expr)

    result = await db.execute(
        base_q
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
        .order_by(sort_expr)
        .offset(skip)
        .limit(limit)
    )
    cases = result.scalars().all()
    return CaseListResponse(
        items=[CaseResponse.model_validate(c) for c in cases],
        total=total,
    )


@router.patch("/bulk-update")
@limiter.limit("20/minute")
async def bulk_update_cases(
    request: Request,
    body: dict,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Bulk update cases: set status or archive/unarchive multiple cases at once.
    Body: { case_ids: [uuid], status?: string, archive?: bool }
    """
    case_ids = body.get("case_ids", [])
    new_status = body.get("status")
    archive = body.get("archive")

    if not case_ids:
        return {"updated": 0}

    from pydantic import TypeAdapter
    case_id_uuids = TypeAdapter(list[UUID]).validate_python(case_ids)

    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id.in_(case_id_uuids))
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
    )
    cases = result.scalars().all()

    actor = _user.display_name if isinstance(_user, UserModel) else "System"
    updated = 0
    for case in cases:
        changed = {}
        if new_status and case.status != new_status:
            changed["status"] = {"old": case.status, "new": new_status}
            case.status = new_status
            # completed_at automatisch setzen/löschen bei Statuswechsel
            if new_status == CaseStatus.COMPLETED and case.completed_at is None:
                case.completed_at = datetime.now(timezone.utc)
            elif new_status != CaseStatus.COMPLETED:
                case.completed_at = None
        if archive is True and case.archived_at is None:
            case.archived_at = datetime.now(timezone.utc)
            changed["archived"] = True
        elif archive is False and case.archived_at is not None:
            case.archived_at = None
            changed["unarchived"] = True
        if changed:
            db.add(ActivityLogModel(
                case_id=case.id,
                event_type="case_updated",
                payload={"actor": actor, "changes": changed, "bulk": True},
            ))
            updated += 1

    await db.flush()
    logger.info("Cases bulk updated", extra={"updated_count": updated, "new_status": new_status, "archive": archive})
    return {"updated": updated}


@router.get("/export")
async def export_cases(
    q: str | None = Query(default=None),
    status: str | None = Query(default=None),
    department: str | None = Query(default=None),
    assignee: str | None = Query(default=None),
    created_by: str | None = Query(default=None),
    has_open_findings: bool | None = Query(default=None),
    include_archived: bool = Query(default=False),
    format: Literal["csv"] = Query(default="csv"),
    db: AsyncSession = Depends(get_db),
    _user=require_roles("viewer", "editor", "admin"),
):
    """Export cases list as CSV. Applies same filters as GET /cases."""
    from datetime import date as date_type
    from sqlalchemy import exists as sql_exists

    base_q = select(CaseModel).options(selectinload(CaseModel.findings))
    if q:
        like = f"%{q}%"
        from sqlalchemy import or_
        base_q = base_q.where(
            or_(
                CaseModel.title.ilike(like),
                CaseModel.department.ilike(like),
                CaseModel.case_type.ilike(like),
                CaseModel.created_by.ilike(like),
                CaseModel.assignee.ilike(like),
            )
        )
    if status:
        base_q = base_q.where(CaseModel.status == status)
    if department:
        base_q = base_q.where(CaseModel.department == department)
    if assignee:
        base_q = base_q.where(CaseModel.assignee.ilike(f"%{assignee}%"))
    if created_by:
        base_q = base_q.where(CaseModel.created_by.ilike(f"%{created_by}%"))
    if has_open_findings is True:
        base_q = base_q.where(
            sql_exists(select(DocumentModel.id).where(
                DocumentModel.case_id == CaseModel.id,
            ))
        )
    if not include_archived:
        base_q = base_q.where(CaseModel.archived_at == None)  # noqa: E711

    result = await db.execute(base_q.order_by(CaseModel.updated_at.desc()).limit(5000))
    cases = result.scalars().all()

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "ID", "Titel", "Abteilung", "Vorgangstyp", "Status", "Ersteller",
        "Zugewiesener", "Frist", "Erstellt am", "Aktualisiert am",
        "Offene Befunde", "Archiviert",
    ])
    for c in cases:
        open_findings = sum(1 for f in c.findings if f.status == FindingStatus.OPEN)
        writer.writerow([
            str(c.id), c.title, c.department, c.case_type, c.status,
            c.created_by, c.assignee,
            c.deadline.isoformat() if c.deadline else "",
            c.created_at.strftime("%Y-%m-%d"),
            c.updated_at.strftime("%Y-%m-%d"),
            open_findings,
            "Ja" if c.archived_at else "Nein",
        ])

    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    body_content = "\ufeff" + buf.getvalue()
    logger.info("Cases exported as CSV", extra={"row_count": len(cases)})
    return Response(
        content=body_content.encode("utf-8"),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="Vorgaenge-{date_str}.csv"'},
    )


@router.get("/{case_id}", response_model=CaseResponse, summary="Vorgang abrufen")
async def get_case(
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    """Get a single case by ID."""
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    return CaseResponse.model_validate(case)


@router.get("/{case_id}/activities", response_model=list[ActivityResponse])
async def get_case_activities(
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    """Get activity log for the case (run_checks, finding_status_updated, etc.), sorted by time descending."""
    result = await db.execute(
        select(CaseModel).where(CaseModel.id == case_id)
    )
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Case not found")
    activities_result = await db.execute(
        select(ActivityLogModel)
        .where(ActivityLogModel.case_id == case_id)
        .order_by(ActivityLogModel.created_at.desc())
    )
    activities = activities_result.scalars().all()
    return [ActivityResponse.model_validate(a) for a in activities]


@router.get("/{case_id}/activities/export", summary="Audit-Trail als CSV exportieren")
async def export_case_activities_csv(
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("viewer", "editor", "admin"),
):
    """Exportiert den vollständigen Audit-Trail eines Vorgangs als CSV (für Behördennachweise)."""
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Vorgang nicht gefunden")

    activities_result = await db.execute(
        select(ActivityLogModel)
        .where(ActivityLogModel.case_id == case_id)
        .order_by(ActivityLogModel.created_at.asc())
    )
    activities = activities_result.scalars().all()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["ID", "Ereignis", "Zeitstempel", "Details"])
    for a in activities:
        payload_str = "; ".join(f"{k}={v}" for k, v in (a.payload or {}).items())
        writer.writerow([
            str(a.id),
            a.event_type,
            a.created_at.strftime("%Y-%m-%d %H:%M:%S UTC"),
            payload_str,
        ])

    filename = f"audit-trail-{case_id}.csv"
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.post("", response_model=CaseResponse, status_code=201, summary="Vorgang erstellen")
@limiter.limit("30/minute")
async def create_case(
    request: Request,
    body: CaseCreate,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Create a new case."""
    case = CaseModel(
        title=body.title,
        department=body.department,
        case_type=body.case_type,
        language=body.language,
        created_by=body.created_by,
        assignee=body.assignee,
        status=CaseStatus.INTAKE,
        playbook_version="",
        processing_context=body.processing_context,
        special_category_data=body.special_category_data,
        international_transfer=body.international_transfer,
        auto_run_checks=body.auto_run_checks,
    )
    db.add(case)
    await db.flush()
    # Re-fetch with relationships loaded to avoid lazy load in async context (MissingGreenlet)
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case.id)
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
    )
    case = result.scalar_one()
    return CaseResponse.model_validate(case)


@router.patch("/{case_id}", response_model=CaseResponse, summary="Vorgang aktualisieren")
@limiter.limit("60/minute")
async def update_case(
    request: Request,
    case_id: UUID,
    body: CaseUpdate,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Update a case (partial update). Returns 403 if case is archived."""
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    if case.archived_at is not None:
        raise HTTPException(status_code=403, detail="Archived cases are read-only. Unarchive first.")
    ALLOWED_UPDATE_FIELDS = {
        "title", "department", "case_type", "status", "language",
        "assignee", "playbook_version", "processing_context",
        "special_category_data", "international_transfer", "deadline",
        "auto_run_checks",
    }
    update_data = body.model_dump(exclude_unset=True)
    changed: dict = {}
    for key, value in update_data.items():
        if key not in ALLOWED_UPDATE_FIELDS:
            raise HTTPException(status_code=422, detail=f"Field not updatable: {key}")
        old_value = getattr(case, key, None)
        if old_value != value:
            changed[key] = {"old": str(old_value) if old_value is not None else None, "new": str(value) if value is not None else None}
        setattr(case, key, value)
    # completed_at automatisch setzen/löschen wenn Status auf "completed" geändert wird
    if "status" in update_data:
        if update_data["status"] == CaseStatus.COMPLETED and case.completed_at is None:
            case.completed_at = datetime.now(timezone.utc)
        elif update_data["status"] != CaseStatus.COMPLETED:
            case.completed_at = None
    await db.flush()
    if changed:
        actor = _user.display_name if isinstance(_user, UserModel) else "System"
        activity = ActivityLogModel(
            case_id=case_id,
            event_type="case_updated",
            payload={"actor": actor, "changes": changed},
        )
        db.add(activity)
        # Webhook-Event feuern wenn Status geändert
        if "status" in changed:
            try:
                from app.services.webhook_service import fire_event
                await fire_event(
                    "case_status_changed",
                    {
                        "case_id": str(case_id),
                        "old_status": changed["status"]["old"],
                        "new_status": changed["status"]["new"],
                        "actor": actor,
                    },
                    db,
                )
            except Exception as exc:
                logger.warning("Webhook fire_event failed (non-critical): %s", exc)
    # Re-fetch with relationships loaded to avoid lazy load in async context (MissingGreenlet)
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
    )
    case = result.scalar_one()
    return CaseResponse.model_validate(case)


@router.delete("/{case_id}", status_code=204)
async def delete_case(
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Delete a case and its documents/findings (cascade)."""
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    await asyncio.to_thread(delete_chunks_by_case_id, case_id)
    await db.delete(case)
    await db.flush()
    return None


@router.post("/{case_id}/audit-export", summary="Audit-Paket exportieren (ZIP)")
@limiter.limit("10/minute")
async def export_audit_package(
    request: Request,
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("viewer", "editor", "admin"),
):
    """Erstellt ein ZIP-Audit-Paket mit Befunden, Dokumentliste, Activity-Log und SHA-256-Manifest.

    Geeignet für externe Prüfer und Aufsichtsbehörden. Kein Ursprungsdokument-Inhalt –
    nur strukturierte Metadaten und Prüfergebnisse.
    """
    from app.services.audit_export_service import build_audit_export
    try:
        zip_bytes, filename = await build_audit_export(case_id, db)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc))
    except Exception as exc:
        logger.error("Audit export failed for case %s: %s", case_id, exc)
        raise HTTPException(status_code=500, detail=f"Export fehlgeschlagen: {exc}")
    from urllib.parse import quote as _quote
    _fn_enc = _quote(filename, safe="")
    _fn_ascii = filename.encode("ascii", errors="replace").decode().replace('"', "_")
    return Response(
        content=zip_bytes,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{_fn_ascii}"; filename*=UTF-8\'\'{_fn_enc}'},
    )


@router.post("/{case_id}/clone", response_model=CaseResponse, status_code=201, summary="Vorgang duplizieren")
@limiter.limit("10/minute")
async def clone_case(
    request: Request,
    case_id: UUID,
    body: CaseCloneRequest,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Dupliziert einen Vorgang: kopiert Metadaten, erstellt keinen neuen Dokument-Upload.
    Optional: LLM passt Verarbeitungskontext auf neuen Titel/Abteilung an."""
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    source = result.scalar_one_or_none()
    if not source:
        raise HTTPException(status_code=404, detail="Case not found")

    new_department = body.new_department or source.department
    new_case = CaseModel(
        title=body.new_title,
        department=new_department,
        case_type=source.case_type,
        language=source.language,
        created_by=_user.display_name if isinstance(_user, UserModel) else "",
        assignee=source.assignee,
        status=CaseStatus.INTAKE,
        playbook_version=source.playbook_version,
        processing_context=source.processing_context,
        special_category_data=source.special_category_data,
        international_transfer=source.international_transfer,
        auto_run_checks=source.auto_run_checks,
        retention_months=source.retention_months,
    )
    db.add(new_case)
    await db.flush()

    # Aktivitätslog auf beiden Vorgängen
    activity_source = ActivityLogModel(
        case_id=case_id,
        event_type="case_cloned",
        payload={"new_case_id": str(new_case.id), "new_title": body.new_title},
    )
    activity_new = ActivityLogModel(
        case_id=new_case.id,
        event_type="case_cloned_from",
        payload={"source_case_id": str(case_id), "source_title": source.title},
    )
    db.add(activity_source)
    db.add(activity_new)
    await db.flush()

    # Re-fetch mit Relationships
    new_result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == new_case.id)
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
    )
    new_case = new_result.scalar_one()
    logger.info(
        "Case cloned",
        extra={"source_case_id": str(case_id), "new_case_id": str(new_case.id)},
    )
    return CaseResponse.model_validate(new_case)


@router.post("/{case_id}/archive", response_model=CaseResponse)
@limiter.limit("30/minute")
async def archive_case(
    request: Request,
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Archive a case. Archived cases are read-only and hidden from the default list view."""
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    if case.archived_at is not None:
        raise HTTPException(status_code=409, detail="Case is already archived")
    case.archived_at = datetime.now(timezone.utc)
    await db.flush()
    await db.refresh(case)
    return CaseResponse.model_validate(case)


@router.post("/{case_id}/unarchive", response_model=CaseResponse)
@limiter.limit("30/minute")
async def unarchive_case(
    request: Request,
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Unarchive a previously archived case."""
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    if case.archived_at is None:
        raise HTTPException(status_code=409, detail="Case is not archived")
    case.archived_at = None
    await db.flush()
    await db.refresh(case)
    return CaseResponse.model_validate(case)


def _build_vvt_export_docx(doc_name: str, source_template: str, fields: list) -> bytes:
    """Build DOCX with VVT normalization (Ziel-Template): document name, template, table of fields."""
    doc = DocxDocument()
    doc.add_heading("VVT-Normalisierung (Ziel-Template)", 0)
    doc.add_paragraph()
    doc.add_paragraph(f"Dokument: {doc_name}")
    doc.add_paragraph(f"Erkanntes Template: {source_template or '—'}")
    doc.add_paragraph()
    table = doc.add_table(rows=1 + len(fields), cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Feldname"
    hdr[1].text = "Status"
    hdr[2].text = "Kanonischer Wert"
    hdr[3].text = "Nachweis"
    hdr[4].text = "Hinweis"
    for i, f in enumerate(fields):
        row = table.rows[i + 1].cells
        row[0].text = (f.field_name or "").replace("\r\n", " ").replace("\n", " ")
        row[1].text = (f.status or "").replace("\r\n", " ").replace("\n", " ")
        row[2].text = (f.canonical_value or "").replace("\r\n", " ").replace("\n", " ")
        row[3].text = (f.evidence or "").replace("\r\n", " ").replace("\n", " ")
        row[4].text = (f.finding or "").replace("\r\n", " ").replace("\n", " ")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/{case_id}/vvt-normalization/export", response_class=Response)
async def get_vvt_normalization_export(
    case_id: UUID,
    document_id: UUID | None = Query(None, alias="document_id"),
    format: Literal["csv", "docx"] = Query("csv", alias="format"),
    db: AsyncSession = Depends(get_db),
):
    """
    Export VVT normalization. Query: format=csv (default) or format=docx.
    Uses first VVT document if document_id not given.
    CSV: text/csv. DOCX: VVT Ziel-Template (document name, template, fields table).
    """
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    vvt_docs = [d for d in case.documents if d.type == "vvt"]
    if not vvt_docs:
        raise HTTPException(status_code=404, detail="No VVT document in this case")

    if document_id is not None:
        doc = next((d for d in vvt_docs if d.id == document_id), None)
        if not doc:
            raise HTTPException(
                status_code=404,
                detail="Document not found or is not a VVT document of this case",
            )
    else:
        doc = vvt_docs[0]

    raw_text = doc.content or ""
    if not raw_text.strip():
        raise HTTPException(
            status_code=404,
            detail="VVT-Dokument hat keinen extrahierten Inhalt",
        )

    case_lang = getattr(case, "language", None)
    vvt_fields = get_vvt_field_names(settings)
    extraction = await normalize_vvt(raw_text, language=case_lang, field_names=vvt_fields)
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    if (format or "csv").lower() == "docx":
        content = _build_vvt_export_docx(doc.name, extraction.source_template or "", extraction.fields)
        filename = f"VVT-Ziel-{case_id}-{date_str}.docx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["document_name", "source_template", "field_name", "status", "canonical_value", "evidence", "finding"])
    for f in extraction.fields:
        writer.writerow([
            doc.name,
            extraction.source_template or "",
            f.field_name,
            f.status,
            (f.canonical_value or "").replace("\r\n", " ").replace("\n", " "),
            (f.evidence or "").replace("\r\n", " ").replace("\n", " "),
            (f.finding or "").replace("\r\n", " ").replace("\n", " "),
        ])
    csv_content = buf.getvalue()
    filename = f"VVT-Export-{case_id}-{date_str}.csv"
    body = "\ufeff" + csv_content  # UTF-8 BOM for Excel
    return Response(
        content=body.encode("utf-8"),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{case_id}/vvt-normalization", response_model=VVTNormalizationResponse)
async def get_vvt_normalization(
    case_id: UUID,
    document_id: UUID | None = None,
    db: AsyncSession = Depends(get_db),
):
    """
    Get VVT normalization for the case. Uses the first VVT document if document_id is not given.
    Returns canonical VVT fields and template detection (LLM-based).
    """
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    vvt_docs = [d for d in case.documents if d.type == "vvt"]
    if not vvt_docs:
        return VVTNormalizationResponse(
            document_id=None,
            document_name="",
            source_template="",
            fields=[],
        )

    if document_id is not None:
        doc = next((d for d in vvt_docs if d.id == document_id), None)
        if not doc:
            raise HTTPException(
                status_code=404,
                detail="Document not found or is not a VVT document of this case",
            )
    else:
        doc = vvt_docs[0]

    raw_text = doc.content or ""
    if not raw_text.strip():
        return VVTNormalizationResponse(
            document_id=doc.id,
            document_name=doc.name,
            source_template="",
            fields=[],
        )

    case_lang = getattr(case, "language", None)
    vvt_fields = get_vvt_field_names(settings)
    extraction = await normalize_vvt(raw_text, language=case_lang, field_names=vvt_fields)

    fields = [
        VVTFieldResponse(
            field_name=f.field_name,
            required=True,
            status=f.status,
            source_template=extraction.source_template,
            canonical_value=f.canonical_value,
            evidence=f.evidence,
            finding=f.finding,
        )
        for f in extraction.fields
    ]
    return VVTNormalizationResponse(
        document_id=doc.id,
        document_name=doc.name,
        source_template=extraction.source_template or "Unbekannt",
        fields=fields,
    )


@router.get("/{case_id}/dsb-report")
async def get_dsb_report(
    case_id: UUID,
    format: Literal["json", "markdown"] = Query("markdown", alias="format"),
    db: AsyncSession = Depends(get_db),
):
    """
    Get stored DSB summary report for the case. Returns 404 if no report exists.
    format=json (with report_stale, stale_reason) or format=markdown (download).
    """
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Case not found")
    report_row = await get_latest_report(case_id, db)
    if not report_row:
        raise HTTPException(status_code=404, detail="No report available. Generate a report first.")
    report = _payload_to_report(report_row.payload)
    result = await db.execute(
        select(CaseModel).where(CaseModel.id == case_id).options(selectinload(CaseModel.documents))
    )
    case = result.scalar_one_or_none()
    current_doc_count = len(case.documents) if case else 0
    current_last_run_at = await get_last_run_checks_at(case_id, db)
    report_stale, stale_reason = compute_stale(report_row, current_doc_count, current_last_run_at)
    if format == "json":
        out = report.model_dump(mode="json")
        out["report_stale"] = report_stale
        out["stale_reason"] = stale_reason
        return JSONResponse(content=out)
    md = render_report_markdown(report)
    slug = re.sub(r"[^\w\s-]", "", report.case_title)[:50].strip() or "Report"
    slug = re.sub(r"[-\s]+", "-", slug)
    date_str = report.generated_at.strftime("%Y-%m-%d")
    filename = f"DSB-Report-{slug}-{date_str}.md"
    return Response(
        content=md,
        media_type="text/markdown; charset=utf-8",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
        },
    )


@router.get("/{case_id}/dsb-report/status")
async def get_dsb_report_status(
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    """Return status of report generation for polling: no_report, running, completed, failed."""
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Case not found")
    report_row = await get_latest_report(case_id, db)
    job_result = await db.execute(
        select(DSBReportJobModel)
        .where(DSBReportJobModel.case_id == case_id)
        .order_by(DSBReportJobModel.created_at.desc())
        .limit(1)
    )
    job = job_result.scalar_one_or_none()
    if job and job.status == JobStatus.RUNNING:
        return {"status": JobStatus.RUNNING, "job_id": str(job.id), "error": None}
    if job and job.status == JobStatus.FAILED:
        return {"status": JobStatus.FAILED, "job_id": str(job.id), "error": job.error}
    if report_row:
        return {"status": JobStatus.COMPLETED, "job_id": None, "error": None}
    return {"status": "no_report", "job_id": None, "error": None}


@router.post("/{case_id}/dsb-report/generate")
@limiter.limit("5/minute")
async def generate_dsb_report(
    request: Request,
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Start DSB report generation. Returns 202 when queued (Celery), 200 when run synchronously."""
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Case not found")
    use_async = settings.celery_enabled and bool((settings.celery_broker_url or "").strip())
    if use_async:
        job = DSBReportJobModel(case_id=case_id, status=JobStatus.RUNNING)
        db.add(job)
        await db.flush()
        await db.refresh(job)
        build_dsb_report_task.delay(str(job.id), get_request_id())
        return JSONResponse(
            status_code=202,
            content={"job_id": str(job.id), "status": JobStatus.RUNNING, "message": "Report wird erstellt."},
        )
    report = await build_dsb_report(case_id, db)
    await save_report(case_id, report, db)
    return JSONResponse(
        status_code=200,
        content={"status": "completed", "message": "Report erstellt.", "generated_at": report.generated_at.isoformat()},
    )


@router.get("/{case_id}/annotated-documents", response_model=list[AnnotatedDocumentListItem])
async def list_annotated_documents(
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    """List documents that have findings and can be downloaded as annotated DOCX."""
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Case not found")
    items = await list_annotatable_documents(case_id, db)
    return [
        AnnotatedDocumentListItem(document_id=doc_id, document_name=name, finding_count=count)
        for doc_id, name, count in items
    ]


@router.get("/{case_id}/annotated-documents/{document_id}", response_class=Response)
async def get_annotated_document(
    case_id: UUID,
    document_id: UUID,
    format: Literal["docx", "pdf"] = Query("docx", alias="format"),
    db: AsyncSession = Depends(get_db),
):
    """Download annotated document (content + findings). Query: format=docx (default) or format=pdf."""
    try:
        if (format or "docx").lower() == "pdf":
            content, filename = await build_annotated_pdf(case_id, document_id, db)
            media_type = "application/pdf"
        else:
            content, filename = await build_annotated_docx(case_id, document_id, db)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{case_id}/run-checks/status")
async def get_run_checks_status(
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    """Return the latest run_checks job status for this case (for polling). Includes last_run activity for timeline and documents_changed_since_last_run flag."""
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Case not found")
    job_result = await db.execute(
        select(RunChecksJobModel)
        .where(RunChecksJobModel.case_id == case_id)
        .order_by(RunChecksJobModel.created_at.desc())
        .limit(1)
    )
    job = job_result.scalar_one_or_none()
    activity_result = await db.execute(
        select(ActivityLogModel)
        .where(ActivityLogModel.case_id == case_id, ActivityLogModel.event_type == "run_checks")
        .order_by(ActivityLogModel.created_at.desc())
        .limit(1)
    )
    last_activity = activity_result.scalar_one_or_none()
    last_run = ActivityResponse.model_validate(last_activity) if last_activity else None

    # Determine if any document was re-uploaded after the last run-checks job
    documents_changed = False
    if job and job.status == JobStatus.COMPLETED:
        max_reupload_result = await db.execute(
            select(func.max(DocumentModel.uploaded_at))
            .where(DocumentModel.case_id == case_id, DocumentModel.version > 1)
        )
        max_reupload = max_reupload_result.scalar_one_or_none()
        if max_reupload and max_reupload > job.created_at:
            documents_changed = True

    if not job:
        return {"status": "never_run", "job_id": None, "playbook_name": None, "findings_count": None, "error": None, "last_run": last_run, "documents_changed_since_last_run": False, "checks_total": 0, "checks_done": 0}
    return {
        "status": job.status,
        "job_id": str(job.id),
        "playbook_name": job.playbook_name,
        "findings_count": job.findings_count,
        "error": job.error,
        "last_run": last_run,
        "documents_changed_since_last_run": documents_changed,
        "checks_total": job.checks_total,
        "checks_done": job.checks_done,
    }


@router.get("/{case_id}/run-checks/stream", summary="Run-Checks-Status als SSE-Stream")
async def stream_run_checks_status(
    case_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    """Server-Sent Events stream for run-checks job progress.

    Emits ``data: <json>`` events every 2 s while a job is running.
    Sends a final ``event: done`` when the job reaches completed/failed/never_run,
    then closes the stream.  The client can use ``EventSource`` instead of polling.

    Event payload matches GET /{case_id}/run-checks/status.
    """
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Case not found")

    async def _event_generator():
        poll_interval = 2  # seconds between DB polls
        max_polls = 300    # ~10 min max stream duration
        for _ in range(max_polls):
            # Fresh session per poll to avoid stale reads
            from app.database import async_session_factory
            async with async_session_factory() as poll_db:
                job_result = await poll_db.execute(
                    select(RunChecksJobModel)
                    .where(RunChecksJobModel.case_id == case_id)
                    .order_by(RunChecksJobModel.created_at.desc())
                    .limit(1)
                )
                job = job_result.scalar_one_or_none()
                activity_result = await poll_db.execute(
                    select(ActivityLogModel)
                    .where(ActivityLogModel.case_id == case_id, ActivityLogModel.event_type == "run_checks")
                    .order_by(ActivityLogModel.created_at.desc())
                    .limit(1)
                )
                last_activity = activity_result.scalar_one_or_none()
                last_run = ActivityResponse.model_validate(last_activity) if last_activity else None

                if not job:
                    payload = {"status": "never_run", "job_id": None, "playbook_name": None,
                               "findings_count": None, "error": None, "last_run": last_run,
                               "checks_total": 0, "checks_done": 0}
                else:
                    payload = {
                        "status": job.status,
                        "job_id": str(job.id),
                        "playbook_name": job.playbook_name,
                        "findings_count": job.findings_count,
                        "error": job.error,
                        "last_run": last_run,
                        "checks_total": job.checks_total,
                        "checks_done": job.checks_done,
                    }

                is_terminal = not job or job.status in (JobStatus.COMPLETED, JobStatus.FAILED, "never_run")
                event_type = "done" if is_terminal else "progress"
                yield f"event: {event_type}\ndata: {_json.dumps(payload, default=str)}\n\n"

            if is_terminal:
                return
            await asyncio.sleep(poll_interval)
        # Timeout: send done with last known state
        yield "event: done\ndata: {}\n\n"

    return StreamingResponse(
        _event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",  # nginx: disable proxy buffering
        },
    )


@router.post("/{case_id}/run-checks", summary="Playbook-Prüfungen starten")
@limiter.limit("10/minute")
async def run_checks(
    request: Request,
    case_id: UUID,
    body: RunChecksRequest,
    db: AsyncSession = Depends(get_db),
    _user=require_roles("editor", "admin"),
):
    """Run playbook checks against all case documents. Returns 202 when queued (Celery), 200 with case when run synchronously."""
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    pb_result = await db.execute(select(PlaybookModel).where(PlaybookModel.id == body.playbook_id))
    playbook = pb_result.scalar_one_or_none()
    if not playbook:
        raise HTTPException(status_code=404, detail="Playbook not found")

    raw_checks = playbook.content.get("checks") if isinstance(playbook.content, dict) else []
    if not raw_checks:
        await db.refresh(case)
        return CaseResponse.model_validate(case)

    strategies = body.strategies or ["full_text"]
    use_async = settings.celery_enabled and bool((settings.celery_broker_url or "").strip())

    # Prevent duplicate concurrent runs: return 409 if a job is already running
    running_job_result = await db.execute(
        select(RunChecksJobModel).where(
            RunChecksJobModel.case_id == case_id,
            RunChecksJobModel.playbook_id == body.playbook_id,
            RunChecksJobModel.status == JobStatus.RUNNING,
        )
    )
    if running_job_result.scalar_one_or_none() is not None:
        raise HTTPException(
            status_code=409,
            detail="A check run for this case and playbook is already in progress. Please wait for it to finish.",
        )

    if use_async:
        job = RunChecksJobModel(
            case_id=case_id,
            status=JobStatus.RUNNING,
            playbook_id=playbook.id,
            playbook_name=playbook.name,
            strategies=strategies,
        )
        db.add(job)
        await db.flush()
        await db.refresh(job)
        # Commit FIRST so the job row is visible to the Celery worker's DB session.
        # Previously the task was dispatched before commit, causing a race condition
        # where the worker could not find the job row.
        await db.commit()
        celery_result = run_playbook_checks.delay(str(job.id), get_request_id())
        job.celery_task_id = celery_result.id
        await db.commit()
        logger.info(
            "run_checks: dispatched celery task",
            extra={
                "job_id": str(job.id),
                "case_id": str(case_id),
                "playbook_name": playbook.name,
                "strategies": strategies,
                "celery_task_id": celery_result.id,
            },
        )
        return JSONResponse(
            status_code=202,
            content={
                "job_id": str(job.id),
                "status": JobStatus.RUNNING,
                "message": "Playbook-Checks werden ausgeführt.",
            },
        )

    findings_added, errors, activity_payload = await run_checks_impl(
        db, case_id, body.playbook_id, strategies, skip_resolved=body.skip_resolved
    )
    activity = ActivityLogModel(
        case_id=case_id,
        event_type="run_checks",
        payload=activity_payload,
    )
    db.add(activity)
    await db.flush()
    result2 = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
    )
    case = result2.scalar_one()
    return CaseResponse.model_validate(case)


@router.get("/{case_id}/risk-score", response_model=CaseRiskScoreResponse)
async def get_case_risk_score(
    case_id: UUID,
    limit: int = Query(default=10, ge=1, le=50),
    db: AsyncSession = Depends(get_db),
):
    """Return current risk score and history derived from completed run_checks jobs.

    Score 0 = no open issues, 100 = maximum risk. Formula mirrors the dashboard complianceScore.
    """
    result = await db.execute(select(CaseModel).where(CaseModel.id == case_id))
    if result.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Case not found")

    jobs_result = await db.execute(
        select(RunChecksJobModel)
        .where(RunChecksJobModel.case_id == case_id, RunChecksJobModel.status == JobStatus.COMPLETED)
        .order_by(RunChecksJobModel.created_at.desc())
        .limit(limit)
    )
    jobs = list(jobs_result.scalars().all())

    case_score_cfg = get_risk_config().case_score

    def _score_from_payload(payload: dict | None, findings_count: int) -> tuple[int, int, int, int]:
        """Extract (critical, high, medium, score) from result_payload or defaults.

        Severity weights and max score come from RiskConfig.case_score so that
        each org-profile can tune the risk model without code changes.
        """
        if payload:
            critical = int(payload.get("critical_findings", 0))
            high = int(payload.get("high_findings", 0))
            medium = int(payload.get("medium_findings", 0))
        else:
            critical = high = medium = 0
        weights = case_score_cfg.severity_weights
        penalty = (
            critical * weights.get("critical", 0)
            + high * weights.get("high", 0)
            + medium * weights.get("medium", 0)
        )
        score = min(case_score_cfg.max_score, penalty)
        return critical, high, medium, score

    history: list[CaseRiskScoreHistoryItem] = []
    for job in reversed(jobs):  # chronological order
        critical, high, medium, score = _score_from_payload(job.result_payload, job.findings_count)
        history.append(CaseRiskScoreHistoryItem(
            job_id=job.id,
            created_at=job.created_at,
            score=score,
            findings_count=job.findings_count,
            critical=critical,
            high=high,
            medium=medium,
        ))

    current_score = history[-1].score if history else 0
    return CaseRiskScoreResponse(case_id=case_id, score=current_score, history=history)


@router.get("/{case_id}/similar", response_model=list[CaseSimilarityResult])
async def get_similar_cases(
    case_id: UUID,
    limit: int = Query(default=5, ge=1, le=10),
    db: AsyncSession = Depends(get_db),
):
    """Return cases in the same department+case_type with overlapping open findings (by check_name)."""
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.findings))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    # Open finding check names for the current case
    current_check_names = {f.check_name for f in case.findings if f.status == FindingStatus.OPEN}
    if not current_check_names:
        return []

    candidates_result = await db.execute(
        select(CaseModel)
        .where(
            CaseModel.department == case.department,
            CaseModel.case_type == case.case_type,
            CaseModel.id != case_id,
        )
        .options(selectinload(CaseModel.findings))
        .limit(50)
    )
    candidates = candidates_result.scalars().all()

    scored: list[tuple[float, CaseModel]] = []
    for candidate in candidates:
        candidate_check_names = {f.check_name for f in candidate.findings}
        shared = current_check_names & candidate_check_names
        if not shared:
            continue
        overlap = len(shared) / len(current_check_names)
        scored.append((overlap, candidate))

    scored.sort(key=lambda x: x[0], reverse=True)
    results = []
    for overlap_score, cand in scored[:limit]:
        resolution = {FindingStatus.FIXED: 0, FindingStatus.ACCEPTED: 0, FindingStatus.OVERRULED: 0}
        shared_check_names = sorted(current_check_names & {f.check_name for f in cand.findings})
        for f in cand.findings:
            if f.check_name in shared_check_names and f.status in resolution:
                resolution[f.status] += 1
        results.append(CaseSimilarityResult(
            case_id=cand.id,
            title=cand.title,
            department=cand.department,
            case_type=cand.case_type,
            status=cand.status,
            overlap_score=round(overlap_score, 2),
            shared_check_names=shared_check_names,
            resolution_summary=resolution,
        ))
    return results
