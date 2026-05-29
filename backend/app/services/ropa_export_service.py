"""ROPA-Export gemäß Art. 30 DSGVO.

Wandelt einen Case + normalisierte VVT-Daten in ein behördentaugliches
ROPA-Dokument um (CSV oder DOCX). Der Service ist als pure-as-possible
Wrapper konzipiert: der DB-Lookup für Case + VVT-Doc liegt im Endpoint,
das Rendering hier.
"""
from __future__ import annotations

import csv
import io
import logging
from datetime import UTC, datetime
from typing import Any, Iterable

logger = logging.getLogger(__name__)


# Canonical Art. 30 Abs. 1 (controller) field ordering. Keys mirror the
# canonical VVT field names used elsewhere; ``label`` is what shows up in
# the rendered CSV/DOCX.
_ART30_SECTIONS: list[tuple[str, str]] = [
    ("verantwortlicher", "1. Name und Kontaktdaten des Verantwortlichen"),
    ("kontakt_dsb", "2. Kontaktdaten des Datenschutzbeauftragten"),
    ("zwecke", "3. Zwecke der Verarbeitung"),
    ("rechtsgrundlage", "4. Rechtsgrundlage(n)"),
    ("kategorien_betroffener", "5. Kategorien betroffener Personen"),
    ("kategorien_daten", "6. Kategorien personenbezogener Daten"),
    ("empfaenger", "7. Kategorien von Empfängern"),
    ("drittland_uebermittlung", "8. Übermittlung in Drittländer / int. Organisationen"),
    ("loeschfristen", "9. Geplante Fristen für die Löschung"),
    ("tom_beschreibung", "10. Allgemeine Beschreibung der TOMs (Art. 32)"),
]


def _vvt_field_to_value(fields: Iterable[dict[str, Any]], wanted_aliases: Iterable[str]) -> str:
    """Resolve the first non-empty canonical_value from a list of VVT field aliases.

    Field-name matching is case-insensitive, substring-tolerant — the VVT
    template variants in the wild use slightly different field names
    (e.g. "Verantwortlicher" vs "Name des Verantwortlichen").
    """
    aliases_lc = [a.lower() for a in wanted_aliases]
    for f in fields:
        name = str(f.get("field_name") or "").lower()
        if any(a in name for a in aliases_lc):
            val = f.get("canonical_value")
            if val:
                return str(val).strip()
    return ""


# Map our Art-30 keys → likely VVT field-name fragments. Keeping this
# central lets future profiles override matching without touching renderers.
_ART30_ALIASES: dict[str, list[str]] = {
    "verantwortlicher": ["verantwortlich", "controller"],
    "kontakt_dsb": ["dsb", "datenschutzbeauftrag", "data protection officer", "dpo"],
    "zwecke": ["zweck", "purpose"],
    "rechtsgrundlage": ["rechtsgrundlage", "legal basis", "rechtsgrund"],
    "kategorien_betroffener": ["betroffene", "data subject", "kategorien betroff"],
    "kategorien_daten": ["datenkategor", "personenbezogen", "data categor"],
    "empfaenger": ["empfänger", "empfaenger", "recipient"],
    "drittland_uebermittlung": ["drittland", "third country", "international"],
    "loeschfristen": ["löschfrist", "loeschfrist", "retention", "aufbewahr"],
    "tom_beschreibung": ["tom", "technisch-organisator", "art. 32", "art 32"],
}


def _project_art30_rows(
    case_summary: dict[str, Any],
    vvt_fields: list[dict[str, Any]],
) -> list[dict[str, str]]:
    """Build a list of {label, key, value, source} rows for the ROPA export."""
    rows: list[dict[str, str]] = []
    for key, label in _ART30_SECTIONS:
        aliases = _ART30_ALIASES.get(key, [key])
        value = _vvt_field_to_value(vvt_fields, aliases)
        # Fall back to case metadata for the first three sections — they
        # are often known organisationally before any VVT document exists.
        if not value:
            value = _case_fallback(key, case_summary)
        rows.append({
            "key": key,
            "label": label,
            "value": value or "— (im VVT-Dokument nicht erfasst)",
            "source": "vvt" if value and value != _case_fallback(key, case_summary) else (
                "case" if value else "missing"
            ),
        })
    return rows


def _case_fallback(key: str, summary: dict[str, Any]) -> str:
    """Org-level defaults for sections that can be answered without a VVT doc."""
    if key == "verantwortlicher":
        return str(summary.get("org_name") or "").strip()
    if key == "kontakt_dsb":
        return str(summary.get("dpo_contact") or "").strip()
    if key == "drittland_uebermittlung":
        return "Ja" if summary.get("international_transfer") else "Nein"
    return ""


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def build_ropa_csv(
    *,
    case_summary: dict[str, Any],
    vvt_fields: list[dict[str, Any]],
) -> bytes:
    """Render the ROPA as a UTF-8 BOM CSV (Excel-friendly)."""
    rows = _project_art30_rows(case_summary, vvt_fields)
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Abschnitt", "Inhalt", "Quelle"])
    for r in rows:
        writer.writerow([r["label"], r["value"], r["source"]])
    writer.writerow([])
    writer.writerow(["Vorgang", case_summary.get("title", ""), ""])
    writer.writerow(["Abteilung", case_summary.get("department", ""), ""])
    writer.writerow(["Stand", datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC"), ""])
    return ("﻿" + buf.getvalue()).encode("utf-8")


def build_ropa_docx(
    *,
    case_summary: dict[str, Any],
    vvt_fields: list[dict[str, Any]],
) -> bytes:
    """Render the ROPA as a DOCX. Requires python-docx (already a project dep)."""
    # Imported lazily so a missing python-docx in a test env (no extras
    # installed) doesn't break the ``import ropa_export_service`` chain
    # used by pure tests.
    from docx import Document
    from docx.shared import Pt

    rows = _project_art30_rows(case_summary, vvt_fields)
    doc = Document()
    org = case_summary.get("org_name") or "Datenschutzagent"
    title = doc.add_heading(f"Verzeichnis von Verarbeitungstätigkeiten – {org}", level=0)
    for run in title.runs:
        run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.add_run("Vorgang: ").bold = True
    meta.add_run(str(case_summary.get("title") or "—"))
    meta.add_run("\nAbteilung: ").bold = True
    meta.add_run(str(case_summary.get("department") or "—"))
    meta.add_run("\nStand: ").bold = True
    meta.add_run(datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC"))

    doc.add_paragraph()  # spacing

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Abschnitt nach Art. 30 DSGVO"
    hdr[1].text = "Inhalt"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for r in rows:
        row = table.add_row().cells
        row[0].text = r["label"]
        row[1].text = r["value"]

    # Provenance footer so an auditor sees what came from where.
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(
        "Quellenzuordnung: Werte aus dem VVT-Dokument sind als 'vvt' "
        "markiert; aus dem Vorgang abgeleitete Felder als 'case'; "
        "fehlende Inhalte als 'missing'. Eine ausführliche Tabelle liefert "
        "der CSV-Export."
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = ["build_ropa_csv", "build_ropa_docx"]
