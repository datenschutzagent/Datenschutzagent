"""Findings export: German label maps and the DOCX report builder.

Kept out of the route module so the document layout is testable without HTTP and
the CSV/DOCX exports share one set of labels.
"""

from __future__ import annotations

import io
from datetime import UTC, datetime

from docx import Document as DocxDocument

from app.constants import FindingSeverity, FindingStatus

SEVERITY_LABELS: dict[str, str] = {
    FindingSeverity.CRITICAL: "Kritisch",
    FindingSeverity.HIGH: "Hoch",
    FindingSeverity.MEDIUM: "Mittel",
    FindingSeverity.LOW: "Niedrig",
    FindingSeverity.INFO: "Info",
}
STATUS_LABELS: dict[str, str] = {
    FindingStatus.OPEN: "Offen",
    FindingStatus.ACCEPTED: "Akzeptiert",
    FindingStatus.OVERRULED: "Überfahren",
    FindingStatus.FIXED: "Behoben",
}


def _add_count_table(doc, header: str, labels: dict[str, str], findings, attr: str):
    table = doc.add_table(rows=1 + len(labels), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = header
    table.rows[0].cells[1].text = "Anzahl"
    for i, (key, label) in enumerate(labels.items()):
        row = table.rows[i + 1]
        row.cells[0].text = label
        row.cells[1].text = str(sum(1 for f in findings if getattr(f, attr) == key))
    doc.add_paragraph()


def _add_finding_section(doc, f, docs_by_id: dict) -> None:
    doc.add_heading(f.check_name, level=2)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = "Table Grid"
    document_label = (
        docs_by_id.get(f.document_id, "Vorgangsbezogen")
        if f.document_id
        else "Vorgangsbezogen"
    )
    rows = [
        ("Schweregrad", SEVERITY_LABELS.get(f.severity, f.severity)),
        ("Status", STATUS_LABELS.get(f.status, f.status)),
        ("Kategorie", f.category),
        ("Dokument", document_label),
        ("Strategie", f.source_strategy or ""),
    ]
    for i, (label, value) in enumerate(rows):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
    doc.add_paragraph()
    doc.add_paragraph("Beschreibung:").bold = True
    doc.add_paragraph(f.description)
    if f.evidence:
        doc.add_paragraph("Nachweise:").bold = True
        for ev in f.evidence:
            doc.add_paragraph(f"• {ev}")
    if f.recommendation:
        doc.add_paragraph("Empfehlung:").bold = True
        doc.add_paragraph(f.recommendation)
    doc.add_paragraph()


def build_findings_docx(case_title: str, findings: list, docs_by_id: dict) -> bytes:
    """Build DOCX with structured findings report: title page, summary table, per-finding sections."""
    doc = DocxDocument()
    doc.add_heading("Befunde-Bericht", 0)
    doc.add_paragraph(f"Vorgang: {case_title}")
    doc.add_paragraph(f"Erstellt: {datetime.now(UTC).strftime('%d.%m.%Y')}")
    doc.add_paragraph()

    doc.add_heading("Zusammenfassung", level=1)
    _add_count_table(doc, "Schweregrad", SEVERITY_LABELS, findings, "severity")
    _add_count_table(doc, "Status", STATUS_LABELS, findings, "status")

    doc.add_heading("Einzelbefunde", level=1)
    for f in findings:
        _add_finding_section(doc, f, docs_by_id)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
