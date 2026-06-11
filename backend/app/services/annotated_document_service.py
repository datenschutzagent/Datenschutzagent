"""Generate annotated DOCX and PDF documents: original content plus findings as comments section."""

from __future__ import annotations

import io
import logging
from uuid import UUID

import fitz  # PyMuPDF
from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.models.db import CaseModel, FindingModel

logger = logging.getLogger(__name__)

# Max characters of document content to include in the generated DOCX (avoid huge files)
MAX_CONTENT_CHARS = 100_000


async def list_annotatable_documents(
    case_id: UUID, db: AsyncSession
) -> list[tuple[UUID, str, int]]:
    """
    Return list of (document_id, document_name, finding_count) for documents in this case
    that have at least one finding.
    """
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents))
    )
    case = result.scalar_one_or_none()
    if not case:
        return []

    from sqlalchemy import func

    counts_result = await db.execute(
        select(FindingModel.document_id, func.count(FindingModel.id))
        .where(FindingModel.case_id == case_id, FindingModel.document_id.isnot(None))
        .group_by(FindingModel.document_id)
    )
    count_by_doc = dict(counts_result.all())

    out: list[tuple[UUID, str, int]] = []
    for doc in case.documents:
        if doc.id in count_by_doc:
            out.append((doc.id, doc.name, count_by_doc[doc.id]))
    return out


async def build_annotated_docx(
    case_id: UUID, document_id: UUID, db: AsyncSession
) -> tuple[bytes, str]:
    """
    Build a DOCX containing the document's extracted content and a findings section.
    Returns (docx_bytes, suggested_filename).
    """
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise ValueError("Case not found")

    doc = next((d for d in case.documents if d.id == document_id), None)
    if not doc:
        raise ValueError("Document not found or not in this case")

    findings = [f for f in case.findings if f.document_id == document_id]

    docx_doc = Document()
    docx_doc.add_heading(doc.name, 0)
    docx_doc.add_paragraph("(Automatisch generierte annotierte Version mit Findings)")

    content = (doc.content or "").strip()
    if content:
        truncated = content[:MAX_CONTENT_CHARS]
        if len(content) > MAX_CONTENT_CHARS:
            truncated += "\n\n[... Inhalt gekürzt ...]"
        docx_doc.add_heading("Dokumentinhalt (extrahiert)", level=1)
        docx_doc.add_paragraph(truncated)

    docx_doc.add_heading("Findings (Prüfergebnisse)", level=1)
    for i, f in enumerate(findings, 1):
        docx_doc.add_heading(f"{i}. {f.check_name} ({f.severity})", level=2)
        docx_doc.add_paragraph(f.description)
        if f.evidence:
            p = docx_doc.add_paragraph()
            p.add_run("Belege: ").bold = True
            p.add_run(" ".join(f.evidence[:3]))  # limit evidence lines
        if f.recommendation:
            p = docx_doc.add_paragraph()
            p.add_run("Empfehlung: ").bold = True
            p.add_run(f.recommendation)
        docx_doc.add_paragraph()

    buf = io.BytesIO()
    docx_doc.save(buf)
    buf.seek(0)

    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in doc.name)
    if not safe_name.rstrip("._"):
        safe_name = "document"
    if not safe_name.endswith(".docx"):
        safe_name += ".docx"
    filename = f"Annotiert-{safe_name}"
    logger.info(
        "Annotated DOCX generated",
        extra={
            "case_id": str(case_id),
            "document_id": str(document_id),
            "findings_count": len(findings),
        },
    )
    return buf.getvalue(), filename


# PDF layout constants
PDF_MARGIN = 50
PDF_LINE_HEIGHT = 14
PDF_FONTSIZE = 11
PDF_PAGE_HEIGHT = 842  # A4 approx
PDF_PAGE_WIDTH = 595
PDF_TEXT_WIDTH = PDF_PAGE_WIDTH - 2 * PDF_MARGIN


def _wrap_lines(text: str, max_chars_per_line: int = 85) -> list[str]:
    """Simple wrap: split by newlines, then break long lines."""
    lines: list[str] = []
    for part in text.replace("\r\n", "\n").split("\n"):
        while len(part) > max_chars_per_line:
            lines.append(part[:max_chars_per_line])
            part = part[max_chars_per_line:]
        if part:
            lines.append(part)
    return lines


async def build_annotated_pdf(
    case_id: UUID, document_id: UUID, db: AsyncSession
) -> tuple[bytes, str]:
    """
    Build a PDF containing the document's extracted content and a findings section.
    Returns (pdf_bytes, suggested_filename).
    """
    result = await db.execute(
        select(CaseModel)
        .where(CaseModel.id == case_id)
        .options(selectinload(CaseModel.documents), selectinload(CaseModel.findings))
    )
    case = result.scalar_one_or_none()
    if not case:
        raise ValueError("Case not found")

    doc = next((d for d in case.documents if d.id == document_id), None)
    if not doc:
        raise ValueError("Document not found or not in this case")

    findings = [f for f in case.findings if f.document_id == document_id]

    pdf_doc = fitz.open()
    y = PDF_MARGIN
    page = pdf_doc.new_page(width=PDF_PAGE_WIDTH, height=PDF_PAGE_HEIGHT)

    def add_block(lines: list[str], fontsize: int = PDF_FONTSIZE) -> None:
        nonlocal y, page
        for line in lines:
            if y > PDF_PAGE_HEIGHT - PDF_MARGIN - PDF_LINE_HEIGHT:
                page = pdf_doc.new_page(width=PDF_PAGE_WIDTH, height=PDF_PAGE_HEIGHT)
                y = PDF_MARGIN
            page.insert_text(
                (PDF_MARGIN, y),
                line[:120],
                fontsize=fontsize,
                fontname="helv",
            )
            y += PDF_LINE_HEIGHT

    add_block([doc.name], fontsize=14)
    add_block(["(Automatisch generierte annotierte Version mit Findings)"])
    y += PDF_LINE_HEIGHT

    content = (doc.content or "").strip()
    if content:
        truncated = content[:MAX_CONTENT_CHARS]
        if len(content) > MAX_CONTENT_CHARS:
            truncated += "\n\n[... Inhalt gekürzt ...]"
        add_block(["Dokumentinhalt (extrahiert)"])
        add_block(_wrap_lines(truncated))
        y += PDF_LINE_HEIGHT

    add_block(["Findings (Prüfergebnisse)"])
    y += PDF_LINE_HEIGHT
    for i, f in enumerate(findings, 1):
        add_block([f"{i}. {f.check_name} ({f.severity})"], fontsize=12)
        add_block(_wrap_lines(f.description or ""))
        if f.evidence:
            add_block(["Belege: " + " ".join(f.evidence[:3])])
        if f.recommendation:
            add_block(["Empfehlung: " + f.recommendation])
        y += PDF_LINE_HEIGHT

    buf = io.BytesIO()
    pdf_doc.save(buf, garbage=4, deflate=True)
    pdf_doc.close()
    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in doc.name)
    if not safe_name.rstrip("._"):
        safe_name = "document"
    if not safe_name.endswith(".pdf"):
        safe_name += ".pdf"
    filename = f"Annotiert-{safe_name}"
    logger.info(
        "Annotated PDF generated",
        extra={
            "case_id": str(case_id),
            "document_id": str(document_id),
            "findings_count": len(findings),
        },
    )
    return buf.getvalue(), filename
