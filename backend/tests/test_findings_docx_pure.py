"""Pure unit tests for the findings DOCX export builder (no DB, no LLM)."""

import io
from types import SimpleNamespace

from docx import Document

from app.constants import FindingStatus
from app.services.findings_export_service import build_findings_docx


def _finding(status: str, severity: str = "high") -> SimpleNamespace:
    return SimpleNamespace(
        check_name="Rechtsgrundlage",
        severity=severity,
        status=status,
        category="Art. 6",
        document_id=None,
        source_strategy="full_text",
        description="desc",
        evidence=["Seite 1"],
        recommendation="fix",
    )


def _table_rows(doc) -> list[list[str]]:
    rows: list[list[str]] = []
    for table in doc.tables:
        for row in table.rows:
            rows.append([c.text for c in row.cells])
    return rows


def test_summary_contains_status_counts():
    """Regression: the status counter was computed and silently discarded."""
    findings = [
        _finding(FindingStatus.OPEN),
        _finding(FindingStatus.OPEN),
        _finding(FindingStatus.FIXED, severity="low"),
    ]
    data = build_findings_docx("Vorgang X", findings, {})
    rows = _table_rows(Document(io.BytesIO(data)))

    assert ["Status", "Anzahl"] in rows
    assert ["Offen", "2"] in rows
    assert ["Behoben", "1"] in rows
    assert ["Akzeptiert", "0"] in rows
    # Severity summary is still present.
    assert ["Hoch", "2"] in rows
    assert ["Niedrig", "1"] in rows


def test_docx_lists_each_finding():
    data = build_findings_docx("Vorgang X", [_finding(FindingStatus.OPEN)], {})
    doc = Document(io.BytesIO(data))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Rechtsgrundlage" in headings
