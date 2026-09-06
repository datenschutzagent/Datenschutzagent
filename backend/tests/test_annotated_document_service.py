"""Tests for app.services.annotated_document_service.

- Pure unit tests for the line wrapper (no DB).
- Service tests that build annotated DOCX/PDF files from real rows written through
  ``async_session_factory`` (require DATABASE_URL).
"""

from __future__ import annotations

import io
import uuid

import fitz  # PyMuPDF
import pytest
from docx import Document

from app.services.annotated_document_service import (
    MAX_CONTENT_CHARS,
    _wrap_lines,
    build_annotated_docx,
    build_annotated_pdf,
    list_annotatable_documents,
)

# asyncio_mode=auto (pytest.ini) runs async tests without an explicit marker;
# the sync unit tests below stay marker-free, the DB-bound ones carry _requires_db.
_requires_db = pytest.mark.requires_db  # skipped by conftest without DATABASE_URL


# ---------------------------------------------------------------------------
# Pure unit tests — _wrap_lines
# ---------------------------------------------------------------------------


def test_wrap_lines_splits_on_newlines_and_drops_empty_parts():
    assert _wrap_lines("a\nb\r\n\nc") == ["a", "b", "c"]


def test_wrap_lines_breaks_long_lines_at_max_chars():
    text = "x" * 200
    lines = _wrap_lines(text, max_chars_per_line=85)
    assert lines == ["x" * 85, "x" * 85, "x" * 30]


def test_wrap_lines_exact_width_is_not_split():
    assert _wrap_lines("y" * 10, max_chars_per_line=10) == ["y" * 10]


def test_wrap_lines_empty_text_yields_no_lines():
    assert _wrap_lines("") == []


# ---------------------------------------------------------------------------
# DB helpers
# ---------------------------------------------------------------------------


async def _create_case(
    *,
    documents: list[dict] | None = None,
    findings: list[dict] | None = None,
) -> tuple[uuid.UUID, list[uuid.UUID]]:
    """Insert a case with documents (returned ids in order) and findings.

    Finding dicts may reference a document by ``doc_index`` (index into ``documents``).
    """
    from app.database import async_session_factory
    from app.models.db import CaseModel, DocumentModel, FindingModel

    case_id = uuid.uuid4()
    doc_ids: list[uuid.UUID] = []
    async with async_session_factory() as session:
        session.add(
            CaseModel(
                id=case_id,
                title=f"Annotated-Doc-Test {uuid.uuid4()}",
                department="IT",
                case_type="Test",
            )
        )
        for i, d in enumerate(documents or []):
            doc_id = uuid.uuid4()
            doc_ids.append(doc_id)
            session.add(
                DocumentModel(
                    id=doc_id,
                    case_id=case_id,
                    name=d.get("name", f"doc-{i}.pdf"),
                    type="other",
                    version=i + 1,  # (case_id, type, version) is UNIQUE
                    format="pdf",
                    size_bytes=10,
                    storage_path=f"test/{case_id}/{i}",
                    content=d.get("content"),
                )
            )
        for i, f in enumerate(findings or []):
            doc_index = f.get("doc_index")
            session.add(
                FindingModel(
                    case_id=case_id,
                    document_id=doc_ids[doc_index] if doc_index is not None else None,
                    check_name=f.get("check_name", f"check-{i}"),
                    severity=f.get("severity", "low"),
                    category="test",
                    description=f.get("description", f"desc-{i}"),
                    evidence=f.get("evidence", []),
                    recommendation=f.get("recommendation", ""),
                )
            )
        await session.commit()
    return case_id, doc_ids


def _docx_paragraphs(data: bytes) -> list[str]:
    return [p.text for p in Document(io.BytesIO(data)).paragraphs]


def _docx_headings(data: bytes) -> list[str]:
    doc = Document(io.BytesIO(data))
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


def _pdf_text(data: bytes) -> tuple[int, str]:
    """Return (page_count, concatenated text) of a PDF."""
    with fitz.open(stream=data, filetype="pdf") as pdf:
        return pdf.page_count, "\n".join(page.get_text() for page in pdf)


# ---------------------------------------------------------------------------
# list_annotatable_documents
# ---------------------------------------------------------------------------


@_requires_db
async def test_list_annotatable_documents_unknown_case_returns_empty():
    from app.database import async_session_factory

    async with async_session_factory() as session:
        assert await list_annotatable_documents(uuid.uuid4(), session) == []


@_requires_db
async def test_list_annotatable_documents_counts_findings_per_document():
    case_id, (doc_a, doc_b, doc_c) = await _create_case(
        documents=[{"name": "a.pdf"}, {"name": "b.pdf"}, {"name": "c.pdf"}],
        findings=[
            {"doc_index": 0},
            {"doc_index": 0},
            {"doc_index": 1},
            {},  # case-level finding without document -> not counted
        ],
    )
    from app.database import async_session_factory

    async with async_session_factory() as session:
        out = await list_annotatable_documents(case_id, session)

    assert {doc_id: (name, n) for doc_id, name, n in out} == {
        doc_a: ("a.pdf", 2),
        doc_b: ("b.pdf", 1),
    }
    # Document without findings is not annotatable.
    assert doc_c not in {doc_id for doc_id, _, _ in out}


@_requires_db
async def test_list_annotatable_documents_case_without_findings_is_empty():
    case_id, _ = await _create_case(documents=[{"name": "only.pdf"}])
    from app.database import async_session_factory

    async with async_session_factory() as session:
        assert await list_annotatable_documents(case_id, session) == []


# ---------------------------------------------------------------------------
# build_annotated_docx
# ---------------------------------------------------------------------------


@_requires_db
async def test_docx_unknown_case_raises():
    from app.database import async_session_factory

    async with async_session_factory() as session:
        with pytest.raises(ValueError, match="Case not found"):
            await build_annotated_docx(uuid.uuid4(), uuid.uuid4(), session)


@_requires_db
async def test_docx_document_not_in_case_raises():
    case_id, _ = await _create_case(documents=[{"name": "x.pdf"}])
    other_case, (foreign_doc,) = await _create_case(documents=[{"name": "y.pdf"}])
    from app.database import async_session_factory

    async with async_session_factory() as session:
        with pytest.raises(ValueError, match="Document not found"):
            await build_annotated_docx(case_id, uuid.uuid4(), session)
        # A document of another case must not be reachable either.
        with pytest.raises(ValueError, match="not in this case"):
            await build_annotated_docx(case_id, foreign_doc, session)


@_requires_db
async def test_docx_contains_content_and_findings():
    case_id, (doc_id, other_doc) = await _create_case(
        documents=[
            {"name": "Vertrag 2026.pdf", "content": "  Präambel des Vertrags.  "},
            {"name": "other.pdf", "content": "irrelevant"},
        ],
        findings=[
            {
                "doc_index": 0,
                "check_name": "Rechtsgrundlage",
                "severity": "high",
                "description": "Keine Rechtsgrundlage genannt.",
                "evidence": ["Zitat 1", "Zitat 2", "Zitat 3", "Zitat 4"],
                "recommendation": "Art. 6 ergänzen.",
            },
            {"doc_index": 1, "check_name": "Fremd-Finding"},
            {"check_name": "Case-Finding"},
        ],
    )
    from app.database import async_session_factory

    async with async_session_factory() as session:
        data, filename = await build_annotated_docx(case_id, doc_id, session)

    assert filename == "Annotiert-Vertrag_2026.pdf.docx"
    assert data[:2] == b"PK"  # zip container

    docx = Document(io.BytesIO(data))
    titles = [p.text for p in docx.paragraphs if p.style.name == "Title"]
    assert titles == ["Vertrag 2026.pdf"]  # add_heading(level=0) -> Title style

    headings = _docx_headings(data)
    assert "Dokumentinhalt (extrahiert)" in headings
    assert "Findings (Prüfergebnisse)" in headings
    assert "1. Rechtsgrundlage (high)" in headings
    # Findings of other documents / case-level findings are excluded.
    assert not any("Fremd-Finding" in h or "Case-Finding" in h for h in headings)

    paragraphs = _docx_paragraphs(data)
    assert "(Automatisch generierte annotierte Version mit Findings)" in paragraphs
    assert "Präambel des Vertrags." in paragraphs  # stripped
    assert "Keine Rechtsgrundlage genannt." in paragraphs
    assert "Belege: Zitat 1 Zitat 2 Zitat 3" in paragraphs  # capped at 3 evidence lines
    assert not any("Zitat 4" in p for p in paragraphs)
    assert "Empfehlung: Art. 6 ergänzen." in paragraphs


@_requires_db
async def test_docx_document_without_content_and_findings_without_evidence():
    case_id, (doc_id,) = await _create_case(
        documents=[{"name": "leer.docx", "content": None}],
        findings=[
            {
                "doc_index": 0,
                "check_name": "Ohne Belege",
                "severity": "medium",
                "description": "Nur Beschreibung.",
                "evidence": [],
                "recommendation": "",
            }
        ],
    )
    from app.database import async_session_factory

    async with async_session_factory() as session:
        data, filename = await build_annotated_docx(case_id, doc_id, session)

    # .docx suffix is not doubled.
    assert filename == "Annotiert-leer.docx"
    headings = _docx_headings(data)
    assert "Dokumentinhalt (extrahiert)" not in headings
    assert "1. Ohne Belege (medium)" in headings
    paragraphs = _docx_paragraphs(data)
    assert "Nur Beschreibung." in paragraphs
    assert not any(p.startswith(("Belege:", "Empfehlung:")) for p in paragraphs)


@_requires_db
async def test_docx_whitespace_only_content_is_treated_as_empty():
    case_id, (doc_id,) = await _create_case(
        documents=[{"name": "blank.pdf", "content": "   \n\t  "}]
    )
    from app.database import async_session_factory

    async with async_session_factory() as session:
        data, _ = await build_annotated_docx(case_id, doc_id, session)

    assert "Dokumentinhalt (extrahiert)" not in _docx_headings(data)
    # No findings -> only the findings heading, no numbered sub-headings.
    assert not any(h.startswith("1.") for h in _docx_headings(data))


@_requires_db
async def test_docx_truncates_oversized_content():
    content = "A" * (MAX_CONTENT_CHARS + 500)
    case_id, (doc_id,) = await _create_case(
        documents=[{"name": "big.pdf", "content": content}]
    )
    from app.database import async_session_factory

    async with async_session_factory() as session:
        data, _ = await build_annotated_docx(case_id, doc_id, session)

    paragraphs = _docx_paragraphs(data)
    body = next(p for p in paragraphs if p.startswith("A" * 100))
    assert body.endswith("[... Inhalt gekürzt ...]")
    assert body.count("A") == MAX_CONTENT_CHARS


@_requires_db
async def test_docx_filename_falls_back_when_name_has_no_safe_chars():
    case_id, (doc_id,) = await _create_case(documents=[{"name": "???"}])
    from app.database import async_session_factory

    async with async_session_factory() as session:
        _, filename = await build_annotated_docx(case_id, doc_id, session)

    assert filename == "Annotiert-document.docx"


# ---------------------------------------------------------------------------
# build_annotated_pdf
# ---------------------------------------------------------------------------


@_requires_db
async def test_pdf_unknown_case_raises():
    from app.database import async_session_factory

    async with async_session_factory() as session:
        with pytest.raises(ValueError, match="Case not found"):
            await build_annotated_pdf(uuid.uuid4(), uuid.uuid4(), session)


@_requires_db
async def test_pdf_document_not_in_case_raises():
    case_id, _ = await _create_case(documents=[{"name": "x.pdf"}])
    from app.database import async_session_factory

    async with async_session_factory() as session:
        with pytest.raises(ValueError, match="Document not found"):
            await build_annotated_pdf(case_id, uuid.uuid4(), session)


@_requires_db
async def test_pdf_contains_content_and_findings():
    case_id, (doc_id, _) = await _create_case(
        documents=[
            {"name": "Vertrag 2026.docx", "content": "Erste Zeile.\nZweite Zeile."},
            {"name": "other.pdf", "content": "irrelevant"},
        ],
        findings=[
            {
                "doc_index": 0,
                "check_name": "Rechtsgrundlage",
                "severity": "high",
                "description": "Keine Rechtsgrundlage genannt.",
                "evidence": ["Zitat 1", "Zitat 2", "Zitat 3", "Zitat 4"],
                "recommendation": "Art. 6 ergaenzen.",
            },
            {"doc_index": 1, "check_name": "Fremd-Finding"},
        ],
    )
    from app.database import async_session_factory

    async with async_session_factory() as session:
        data, filename = await build_annotated_pdf(case_id, doc_id, session)

    assert filename == "Annotiert-Vertrag_2026.docx.pdf"
    assert data.startswith(b"%PDF")

    pages, text = _pdf_text(data)
    assert pages == 1
    assert "Vertrag 2026.docx" in text
    assert "(Automatisch generierte annotierte Version mit Findings)" in text
    assert "Dokumentinhalt (extrahiert)" in text
    assert "Erste Zeile." in text
    assert "Zweite Zeile." in text
    assert "1. Rechtsgrundlage (high)" in text
    assert "Keine Rechtsgrundlage genannt." in text
    assert "Belege: Zitat 1 Zitat 2 Zitat 3" in text
    assert "Zitat 4" not in text
    assert "Empfehlung: Art. 6 ergaenzen." in text
    assert "Fremd-Finding" not in text


@_requires_db
async def test_pdf_without_content_and_without_evidence():
    case_id, (doc_id,) = await _create_case(
        documents=[{"name": "leer.pdf", "content": None}],
        findings=[
            {
                "doc_index": 0,
                "check_name": "Ohne Belege",
                "severity": "medium",
                "description": "Nur Beschreibung.",
                "evidence": [],
                "recommendation": "",
            }
        ],
    )
    from app.database import async_session_factory

    async with async_session_factory() as session:
        data, filename = await build_annotated_pdf(case_id, doc_id, session)

    assert filename == "Annotiert-leer.pdf"  # suffix not doubled
    pages, text = _pdf_text(data)
    assert pages == 1
    assert "Dokumentinhalt (extrahiert)" not in text
    assert "1. Ohne Belege (medium)" in text
    assert "Nur Beschreibung." in text
    assert "Belege:" not in text
    assert "Empfehlung:" not in text


@_requires_db
async def test_pdf_long_content_spans_multiple_pages():
    # 200 lines at 14pt line height exceed one A4 page (~53 usable lines).
    content = "\n".join(f"Zeile {i:03d}" for i in range(200))
    case_id, (doc_id,) = await _create_case(
        documents=[{"name": "lang.pdf", "content": content}]
    )
    from app.database import async_session_factory

    async with async_session_factory() as session:
        data, _ = await build_annotated_pdf(case_id, doc_id, session)

    pages, text = _pdf_text(data)
    assert pages >= 3
    assert "Zeile 000" in text
    assert "Zeile 199" in text
    assert "Findings (Prüfergebnisse)" in text


@_requires_db
async def test_pdf_truncates_oversized_content():
    content = "B" * (MAX_CONTENT_CHARS + 500)
    case_id, (doc_id,) = await _create_case(
        documents=[{"name": "big.pdf", "content": content}]
    )
    from app.database import async_session_factory

    async with async_session_factory() as session:
        data, _ = await build_annotated_pdf(case_id, doc_id, session)

    pages, text = _pdf_text(data)
    assert "[... Inhalt gekürzt ...]" in text
    # 100_000 chars wrapped at 85 chars/line -> ~1177 lines -> many pages.
    assert pages > 20


@_requires_db
async def test_pdf_filename_falls_back_when_name_has_no_safe_chars():
    case_id, (doc_id,) = await _create_case(documents=[{"name": "..."}])
    from app.database import async_session_factory

    async with async_session_factory() as session:
        _, filename = await build_annotated_pdf(case_id, doc_id, session)

    assert filename == "Annotiert-document.pdf"
