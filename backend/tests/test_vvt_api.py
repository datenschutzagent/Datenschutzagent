"""Integration tests for the /api/v1/vvt-overview routes.

These tests require a live PostgreSQL database (DATABASE_URL env var).
They test VVT list, stats, and export endpoints.
"""

import pytest

from tests.factories import create_case

pytestmark = pytest.mark.asyncio


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# GET /api/v1/vvt-overview
# ---------------------------------------------------------------------------


async def test_vvt_overview_returns_200(client):
    resp = await client.get("/api/v1/vvt-overview")
    assert resp.status_code == 200


async def test_vvt_overview_returns_list(client):
    resp = await client.get("/api/v1/vvt-overview")
    assert isinstance(resp.json(), list)


async def test_vvt_overview_contains_created_case(client):
    case = await create_case(client, title="VVT Overview Test")
    resp = await client.get("/api/v1/vvt-overview")
    assert resp.status_code == 200
    case_ids = [item["case_id"] for item in resp.json()]
    assert case["id"] in case_ids


async def test_vvt_overview_item_has_required_fields(client):
    await create_case(client, title="VVT Field Test")
    resp = await client.get("/api/v1/vvt-overview")
    assert resp.status_code == 200
    items = resp.json()
    if items:
        item = items[0]
        assert "case_id" in item
        assert "title" in item
        assert "department" in item
        assert "case_type" in item
        assert "status" in item
        assert "has_vvt_document" in item


async def test_vvt_overview_new_case_has_no_vvt(client):
    case = await create_case(client, title="No VVT Test")
    resp = await client.get("/api/v1/vvt-overview")
    items = {i["case_id"]: i for i in resp.json()}
    if case["id"] in items:
        assert items[case["id"]]["has_vvt_document"] is False


async def test_vvt_overview_department_filter(client):
    dept = "Spezial-VVT-Abteilung"
    case = await create_case(client, title="Dept Filter Test", department=dept)
    resp = await client.get("/api/v1/vvt-overview", params={"department": dept})
    assert resp.status_code == 200
    ids = [i["case_id"] for i in resp.json()]
    assert case["id"] in ids


async def test_vvt_overview_filter_has_vvt_false(client):
    """Filter has_vvt=false returns only cases without VVT documents."""
    case = await create_case(client, title="No VVT Filter Test")
    resp = await client.get("/api/v1/vvt-overview", params={"has_vvt": "false"})
    assert resp.status_code == 200
    ids = [i["case_id"] for i in resp.json()]
    assert case["id"] in ids


async def test_vvt_overview_pagination(client):
    resp = await client.get("/api/v1/vvt-overview", params={"skip": 0, "limit": 5})
    assert resp.status_code == 200
    assert len(resp.json()) <= 5


# ---------------------------------------------------------------------------
# GET /api/v1/vvt-overview/stats
# ---------------------------------------------------------------------------


async def test_vvt_stats_returns_200(client):
    resp = await client.get("/api/v1/vvt-overview/stats")
    assert resp.status_code == 200


async def test_vvt_stats_has_required_fields(client):
    resp = await client.get("/api/v1/vvt-overview/stats")
    body = resp.json()
    assert "total_cases" in body
    assert "with_vvt" in body
    assert "without_vvt" in body
    assert "by_department" in body
    assert "by_case_type" in body


async def test_vvt_stats_counts_are_non_negative(client):
    resp = await client.get("/api/v1/vvt-overview/stats")
    body = resp.json()
    assert body["total_cases"] >= 0
    assert body["with_vvt"] >= 0
    assert body["without_vvt"] >= 0


# ---------------------------------------------------------------------------
# GET /api/v1/vvt-overview/export
# ---------------------------------------------------------------------------


async def test_vvt_overview_export_csv_returns_200(client):
    resp = await client.get("/api/v1/vvt-overview/export", params={"format": "csv"})
    assert resp.status_code == 200
    assert "csv" in resp.headers.get("content-type", "").lower()


async def test_vvt_overview_export_unsupported_format_returns_400(client):
    resp = await client.get("/api/v1/vvt-overview/export", params={"format": "xlsx"})
    assert resp.status_code == 400


# ---------------------------------------------------------------------------
# Case-level VVT normalization: /api/v1/cases/{case_id}/vvt-normalization[/export]
# (app/api/routes/cases/vvt.py). The LLM-backed ``normalize_vvt`` is patched in
# the route module so no provider is required.
# ---------------------------------------------------------------------------

NIL_UUID = "00000000-0000-0000-0000-000000000000"


async def _add_vvt_document(
    case_id: str,
    *,
    content: str | None,
    doc_type: str = "vvt",
    version: int = 1,
    name: str | None = None,
) -> str:
    """Insert a document row directly (bypasses upload/MinIO) and return its id."""
    import uuid

    from app.database import async_session_factory
    from app.models.db import DocumentModel

    async with async_session_factory() as session:
        doc = DocumentModel(
            case_id=uuid.UUID(case_id),
            name=name or f"{doc_type}-v{version}.xlsx",
            type=doc_type,
            version=version,
            format="xlsx",
            size_bytes=42,
            storage_path=f"test/{case_id}/{doc_type}-v{version}.xlsx",
            content=content,
            extraction_status="done",
        )
        session.add(doc)
        await session.commit()
        return str(doc.id)


def _fake_extraction(source_template: str = "Variante A"):
    """Deterministic stand-in for the LLM extraction result."""
    from app.services.vvt_service import _VVTExtractionField, _VVTExtractionResult

    return _VVTExtractionResult(
        source_template=source_template,
        fields=[
            _VVTExtractionField(
                field_name="Verantwortlicher",
                status="filled",
                canonical_value="Musterstadt GmbH",
                evidence="Sheet 1, Zeile 2",
                finding=None,
            ),
            _VVTExtractionField(
                field_name="Löschfristen",
                status="inconsistent",
                canonical_value="3 Jahre\nbzw. 5 Jahre",
                evidence="Sheet 1,\r\nZeile 9",
                finding="Widerspruch\nzwischen Zeilen",
            ),
            _VVTExtractionField(
                field_name="Empfänger",
                status="missing",
                canonical_value=None,
                evidence=None,
                finding=None,
            ),
        ],
    )


def _patch_normalize(monkeypatch, extraction):
    from unittest.mock import AsyncMock

    from app.api.routes.cases import vvt as vvt_module

    mock = AsyncMock(return_value=extraction)
    monkeypatch.setattr(vvt_module, "normalize_vvt", mock)
    return mock


# ---- GET /cases/{case_id}/vvt-normalization ------------------------------


async def test_vvt_normalization_unknown_case_returns_404(client):
    resp = await client.get(f"/api/v1/cases/{NIL_UUID}/vvt-normalization")
    assert resp.status_code == 404
    assert resp.json()["detail"] == "Case not found"


async def test_vvt_normalization_case_without_vvt_document_is_empty(
    client, monkeypatch
):
    mock = _patch_normalize(monkeypatch, _fake_extraction())
    case = await create_case(client, title="VVT-Norm ohne Dokument")
    # a non-VVT document must not be picked up
    await _add_vvt_document(case["id"], content="irrelevant", doc_type="avv")

    resp = await client.get(f"/api/v1/cases/{case['id']}/vvt-normalization")
    assert resp.status_code == 200
    body = resp.json()
    assert body["document_id"] is None
    assert body["document_name"] == ""
    assert body["source_template"] == ""
    assert body["fields"] == []
    mock.assert_not_awaited()


async def test_vvt_normalization_document_without_content_returns_no_fields(
    client, monkeypatch
):
    mock = _patch_normalize(monkeypatch, _fake_extraction())
    case = await create_case(client, title="VVT-Norm leerer Inhalt")
    doc_id = await _add_vvt_document(case["id"], content="   \n", name="leer.xlsx")

    resp = await client.get(f"/api/v1/cases/{case['id']}/vvt-normalization")
    assert resp.status_code == 200
    body = resp.json()
    assert body["document_id"] == doc_id
    assert body["document_name"] == "leer.xlsx"
    assert body["source_template"] == ""
    assert body["fields"] == []
    mock.assert_not_awaited()


async def test_vvt_normalization_unknown_document_id_returns_404(client, monkeypatch):
    _patch_normalize(monkeypatch, _fake_extraction())
    case = await create_case(client, title="VVT-Norm falsches Dokument")
    await _add_vvt_document(case["id"], content="Verantwortlicher: Musterstadt GmbH")
    other_doc = await _add_vvt_document(case["id"], content="AVV", doc_type="avv")

    resp = await client.get(
        f"/api/v1/cases/{case['id']}/vvt-normalization",
        params={"document_id": other_doc},
    )
    assert resp.status_code == 404
    assert "not a VVT document" in resp.json()["detail"]

    resp = await client.get(
        f"/api/v1/cases/{case['id']}/vvt-normalization",
        params={"document_id": NIL_UUID},
    )
    assert resp.status_code == 404


async def test_vvt_normalization_returns_mapped_fields(client, monkeypatch):
    mock = _patch_normalize(monkeypatch, _fake_extraction("Variante A"))
    case = await create_case(client, title="VVT-Norm Erfolg", language="de")
    raw = "Verantwortlicher: Musterstadt GmbH\nLöschfristen: 3 Jahre"
    doc_id = await _add_vvt_document(case["id"], content=raw, name="vvt.xlsx")

    resp = await client.get(f"/api/v1/cases/{case['id']}/vvt-normalization")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["document_id"] == doc_id
    assert body["document_name"] == "vvt.xlsx"
    assert body["source_template"] == "Variante A"
    assert [f["field_name"] for f in body["fields"]] == [
        "Verantwortlicher",
        "Löschfristen",
        "Empfänger",
    ]
    filled = body["fields"][0]
    assert filled["status"] == "filled"
    assert filled["required"] is True
    assert filled["source_template"] == "Variante A"
    assert filled["canonical_value"] == "Musterstadt GmbH"
    assert filled["evidence"] == "Sheet 1, Zeile 2"
    assert filled["finding"] is None
    missing = body["fields"][2]
    assert missing["status"] == "missing"
    assert missing["canonical_value"] is None

    mock.assert_awaited_once()
    args, kwargs = mock.await_args
    assert args[0] == raw
    assert kwargs["language"] == "de"
    assert isinstance(kwargs["field_names"], list) and kwargs["field_names"]


async def test_vvt_normalization_unknown_template_falls_back_to_unbekannt(
    client, monkeypatch
):
    _patch_normalize(monkeypatch, _fake_extraction(source_template=""))
    case = await create_case(client, title="VVT-Norm Template unbekannt")
    await _add_vvt_document(case["id"], content="Irgendein Text")

    resp = await client.get(f"/api/v1/cases/{case['id']}/vvt-normalization")
    assert resp.status_code == 200
    assert resp.json()["source_template"] == "Unbekannt"
    # per-field template mirrors the raw (empty) extraction value
    assert all(f["source_template"] == "" for f in resp.json()["fields"])


async def test_vvt_normalization_explicit_document_id_selects_document(
    client, monkeypatch
):
    mock = _patch_normalize(monkeypatch, _fake_extraction())
    case = await create_case(client, title="VVT-Norm zweites Dokument")
    await _add_vvt_document(case["id"], content="Erste Version", version=1)
    second = await _add_vvt_document(
        case["id"], content="Zweite Version", version=2, name="vvt-v2.xlsx"
    )

    resp = await client.get(
        f"/api/v1/cases/{case['id']}/vvt-normalization",
        params={"document_id": second},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["document_id"] == second
    assert resp.json()["document_name"] == "vvt-v2.xlsx"
    assert mock.await_args.args[0] == "Zweite Version"


# ---- GET /cases/{case_id}/vvt-normalization/export -----------------------


async def test_vvt_export_unknown_case_returns_404(client):
    resp = await client.get(f"/api/v1/cases/{NIL_UUID}/vvt-normalization/export")
    assert resp.status_code == 404
    assert resp.json()["detail"] == "Case not found"


async def test_vvt_export_without_vvt_document_returns_404(client, monkeypatch):
    mock = _patch_normalize(monkeypatch, _fake_extraction())
    case = await create_case(client, title="VVT-Export ohne Dokument")
    resp = await client.get(f"/api/v1/cases/{case['id']}/vvt-normalization/export")
    assert resp.status_code == 404
    assert resp.json()["detail"] == "No VVT document in this case"
    mock.assert_not_awaited()


async def test_vvt_export_unknown_document_id_returns_404(client, monkeypatch):
    _patch_normalize(monkeypatch, _fake_extraction())
    case = await create_case(client, title="VVT-Export falsches Dokument")
    await _add_vvt_document(case["id"], content="Text")
    resp = await client.get(
        f"/api/v1/cases/{case['id']}/vvt-normalization/export",
        params={"document_id": NIL_UUID},
    )
    assert resp.status_code == 404
    assert "not a VVT document" in resp.json()["detail"]


async def test_vvt_export_document_without_content_returns_404(client, monkeypatch):
    mock = _patch_normalize(monkeypatch, _fake_extraction())
    case = await create_case(client, title="VVT-Export leerer Inhalt")
    await _add_vvt_document(case["id"], content=None)
    resp = await client.get(f"/api/v1/cases/{case['id']}/vvt-normalization/export")
    assert resp.status_code == 404
    assert "keinen extrahierten Inhalt" in resp.json()["detail"]
    mock.assert_not_awaited()


async def test_vvt_export_invalid_format_returns_422(client, monkeypatch):
    _patch_normalize(monkeypatch, _fake_extraction())
    case = await create_case(client, title="VVT-Export Format")
    await _add_vvt_document(case["id"], content="Text")
    resp = await client.get(
        f"/api/v1/cases/{case['id']}/vvt-normalization/export",
        params={"format": "xlsx"},
    )
    assert resp.status_code == 422


async def test_vvt_export_csv_default(client, monkeypatch):
    import csv
    import io

    mock = _patch_normalize(monkeypatch, _fake_extraction("Variante B"))
    case = await create_case(client, title="VVT-Export CSV", language="en")
    await _add_vvt_document(
        case["id"], content="Controller: Example Ltd", name="r.xlsx"
    )

    resp = await client.get(f"/api/v1/cases/{case['id']}/vvt-normalization/export")
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith("text/csv")
    disposition = resp.headers["content-disposition"]
    assert disposition.startswith("attachment;")
    assert f"VVT-Export-{case['id']}-" in disposition
    assert disposition.endswith('.csv"')

    assert resp.content.startswith(b"\xef\xbb\xbf")  # UTF-8 BOM for Excel
    rows = list(csv.reader(io.StringIO(resp.content.decode("utf-8-sig"))))
    assert rows[0] == [
        "document_name",
        "source_template",
        "field_name",
        "status",
        "canonical_value",
        "evidence",
        "finding",
    ]
    assert len(rows) == 4
    assert rows[1] == [
        "r.xlsx",
        "Variante B",
        "Verantwortlicher",
        "filled",
        "Musterstadt GmbH",
        "Sheet 1, Zeile 2",
        "",
    ]
    # line breaks inside values are flattened to single-line cells
    assert rows[2][2:] == [
        "Löschfristen",
        "inconsistent",
        "3 Jahre bzw. 5 Jahre",
        "Sheet 1, Zeile 9",
        "Widerspruch zwischen Zeilen",
    ]
    assert rows[3][2:] == ["Empfänger", "missing", "", "", ""]
    assert mock.await_args.kwargs["language"] == "en"


async def test_vvt_export_docx(client, monkeypatch):
    import io

    from docx import Document as DocxDocument

    _patch_normalize(monkeypatch, _fake_extraction("Variante A"))
    case = await create_case(client, title="VVT-Export DOCX")
    await _add_vvt_document(case["id"], content="Verantwortlicher: X", name="v.xlsx")

    resp = await client.get(
        f"/api/v1/cases/{case['id']}/vvt-normalization/export",
        params={"format": "docx"},
    )
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    disposition = resp.headers["content-disposition"]
    assert f"VVT-Ziel-{case['id']}-" in disposition
    assert disposition.endswith('.docx"')
    assert resp.content[:2] == b"PK"  # OOXML zip container

    document = DocxDocument(io.BytesIO(resp.content))
    paragraphs = [p.text for p in document.paragraphs]
    assert "VVT-Normalisierung (Ziel-Template)" in paragraphs
    assert "Dokument: v.xlsx" in paragraphs
    assert "Erkanntes Template: Variante A" in paragraphs
    table = document.tables[0]
    assert len(table.rows) == 4  # header + 3 fields
    header = [c.text for c in table.rows[0].cells]
    assert header == ["Feldname", "Status", "Kanonischer Wert", "Nachweis", "Hinweis"]
    assert [c.text for c in table.rows[1].cells] == [
        "Verantwortlicher",
        "filled",
        "Musterstadt GmbH",
        "Sheet 1, Zeile 2",
        "",
    ]
    assert [c.text for c in table.rows[2].cells] == [
        "Löschfristen",
        "inconsistent",
        "3 Jahre bzw. 5 Jahre",
        "Sheet 1, Zeile 9",
        "Widerspruch zwischen Zeilen",
    ]


async def test_vvt_export_docx_without_detected_template_shows_dash(
    client, monkeypatch
):
    import io

    from docx import Document as DocxDocument

    _patch_normalize(monkeypatch, _fake_extraction(source_template=""))
    case = await create_case(client, title="VVT-Export DOCX ohne Template")
    await _add_vvt_document(case["id"], content="Text")

    resp = await client.get(
        f"/api/v1/cases/{case['id']}/vvt-normalization/export",
        params={"format": "docx"},
    )
    assert resp.status_code == 200
    document = DocxDocument(io.BytesIO(resp.content))
    assert "Erkanntes Template: —" in [p.text for p in document.paragraphs]


async def test_vvt_export_explicit_document_id(client, monkeypatch):
    mock = _patch_normalize(monkeypatch, _fake_extraction())
    case = await create_case(client, title="VVT-Export zweites Dokument")
    await _add_vvt_document(case["id"], content="Erste Version", version=1)
    second = await _add_vvt_document(
        case["id"], content="Zweite Version", version=2, name="vvt-v2.xlsx"
    )
    resp = await client.get(
        f"/api/v1/cases/{case['id']}/vvt-normalization/export",
        params={"document_id": second, "format": "csv"},
    )
    assert resp.status_code == 200
    assert mock.await_args.args[0] == "Zweite Version"
    assert "vvt-v2.xlsx" in resp.content.decode("utf-8-sig")
